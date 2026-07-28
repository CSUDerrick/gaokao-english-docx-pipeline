#!/usr/bin/env python3
"""Build the student vocabulary handout: two tables, no answers, no attribution.

Matches the shape of the teacher's own 重难点阅读词汇.docx — a title followed by
one long bordered table — but adds the second table (word formation) that the
语法填空 questions actually test.

The words come from the `vocab` stage, which reads only the question text. The
handout goes to students, so nothing here may reveal which paper a passage came
from, and no answer key may reach it.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx.table import _Cell

import docx_splice as ds
from bundle_paths import template_dir

VOCAB = "高三英语精选试题_重难点词汇表_学生版.docx"

WORDS_HEADER = ("英文单词", "词性", "准确的中文释义")
FORMS_HEADER = ("基础词汇", "词性", "变形后的词汇", "变形后的词性", "考点说明")


def _clean(value: object) -> str:
    return ds.sanitize(str(value or "")).strip()


def collect_words(rows: list[dict]) -> list[tuple[str, str, str]]:
    """Dedupe reading words across papers, keeping first appearance order.

    This is the "拼接" step: each row is one paper's list, and the two handout tables are
    every paper's words merged. Deduping here rather than in the prompt means a word that
    appears in two papers costs nothing extra to notice.
    """
    seen: dict[str, tuple[str, str, str]] = {}
    for row in rows:
        for entry in row.get("reading_words") or []:
            if not isinstance(entry, dict):
                continue
            word = _clean(entry.get("word"))
            if not word or word.lower() in seen:
                continue
            seen[word.lower()] = (word, _clean(entry.get("pos")), _clean(entry.get("meaning")))
    return list(seen.values())


def collect_forms(rows: list[dict]) -> list[tuple[str, str, str, str, str]]:
    """Dedupe word-formation pairs by (base, derived), across papers."""
    seen: dict[tuple[str, str], tuple[str, str, str, str, str]] = {}
    for row in rows:
        for entry in row.get("word_forms") or []:
            if not isinstance(entry, dict):
                continue
            base = _clean(entry.get("base"))
            derived = _clean(entry.get("derived"))
            key = (base.lower(), derived.lower())
            if not base or not derived or key in seen:
                continue
            seen[key] = (
                base,
                _clean(entry.get("base_pos")),
                derived,
                _clean(entry.get("derived_pos")),
                _clean(entry.get("note")),
            )
    return list(seen.values())


def _add_table(doc, header: tuple[str, ...], rows: list[tuple[str, ...]]) -> None:
    """Fill one bordered table, linear in row count.

    Do not index `table.cell(r, col)` here: python-docx rebuilds the whole cell grid on
    every call, so filling the table is quadratic. One paper's ~250 rows hid it; merging
    a term's worth of papers into one list (1000+ rows) spent 60–80s per table. Walking
    each new row's own `tc` elements is the same output — no merged cells in these two
    tables — and finishes instantly.
    """
    table = doc.add_table(rows=0, cols=len(header))
    ds.set_table_borders(table)
    ds.fit_table_to_window(table)
    style = doc.styles[ds.VOCAB_BODY]

    for r, values in enumerate([header, *rows]):
        tr = table.add_row()._tr
        for col, tc in enumerate(tr.tc_lst):
            cell = _Cell(tc, table)
            cell.text = values[col] if col < len(values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.style = style
            for run in paragraph.runs:
                run.bold = r == 0  # header row only


def build(rows: list[dict], out: Path) -> Path:
    doc = ds.blank_template(template_dir() / "student_reference.docx")
    ds.ensure_vocab_styles(doc)
    # The handout is printed and handed out, so: no running header, and the same
    # 0.5" margins the teacher's own word lists use — the words then fit the page
    # the way she already trims them to.
    ds.drop_headers(doc)
    ds.set_page_margins(doc, 720)

    words = collect_words(rows)
    forms = collect_forms(rows)

    doc.add_paragraph("重难点阅读词汇", style=ds.VOCAB_TITLE)
    if words:
        _add_table(doc, WORDS_HEADER, words)
    else:
        doc.add_paragraph("暂无", style=ds.VOCAB_BODY)

    # The two lists are studied and marked up separately, so they must not share a
    # page: the teacher prints one and hands out the other.
    forms_title = doc.add_paragraph("语法词汇变形", style=ds.VOCAB_TITLE)
    ds.page_break_before(forms_title)
    if forms:
        _add_table(doc, FORMS_HEADER, forms)
    else:
        doc.add_paragraph("暂无", style=ds.VOCAB_BODY)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    ds.scrub_metadata(out, "高三英语重难点词汇表")
    ds.validate(out)
    return out


def export_vocab(out_dir: Path, log=print) -> Path | None:
    index = out_dir / "vocab_index.json"
    if not index.exists():
        log("  跳过词汇表：未找到 vocab_index.json（先跑 --mode vocab）")
        return None
    rows = json.loads(index.read_text(encoding="utf-8"))
    if not isinstance(rows, list) or not rows:
        log("  跳过词汇表：vocab_index.json 为空")
        return None

    target = out_dir / "docx_exports" / VOCAB
    build(rows, target)
    words, forms = len(collect_words(rows)), len(collect_forms(rows))
    log(f"  wrote {target.name}  (词汇 {words} 条，词形变化 {forms} 条)")
    return target
