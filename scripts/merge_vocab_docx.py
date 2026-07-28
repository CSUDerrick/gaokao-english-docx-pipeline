#!/usr/bin/env python3
"""Merge word-list handouts that were generated separately into one deduped handout.

A term's worth of papers produces a term's worth of 重难点词汇表 — one per run, and
some runs wrote the two lists (重难点阅读词汇 / 语法词汇变形) into separate files. A
student revising for the exam wants one sheet, with each word on it once.

This is entirely local: it reads the handouts already on disk and calls no API. The
words are never re-generated, only re-grouped, so nothing here can invent a meaning
the model did not produce — and the merged file goes out through the same
``export_vocab_docx.build`` the pipeline uses, so it is byte-for-byte the same shape
as the handouts it was built from (小四/TNR+宋体, bordered tables, the two lists on
separate pages) and passes the same gates.

Dedupe rules, and why they are what they are:

* **Reading words are keyed on the word alone**, not on (word, 词性). Merging papers
  turns up ~9 words per term whose 词性 genuinely differs between papers (``yield``
  n. 产量 / v. 屈服). Keying on the pair would print those twice with no hint that they
  are the same word; keying on the word alone and unioning both is one row that says
  ``n./v.``. It also means ``ripple effect`` tagged ``n. phr.`` in one paper and
  ``phr.`` in another does not become two rows.
* **Meanings are unioned per sense, not per string.** Three papers writing
  「大量，丰富」/「丰富，充裕」/「充裕，大量」 are one word with three senses, and joining
  the strings would print each sense up to three times. So each gloss is split on its
  own separators and the senses are unioned. A sense that carries a 熟词生义 note wins
  over the bare form of the same sense — that note is the only reason the word is on
  the sheet.
* **考点说明 is one explanation, not a list**: variants that differ only in spacing
  (「v.变n.，加后缀 -ment」 vs 「v. 变 n.，加后缀 -ment」) must not be concatenated, so the
  most informative single variant is kept.

Both tables come out in alphabetical order. First-appearance order is what the
per-paper handout uses, and it is the right order there (it follows the paper); across
a whole term it is no order at all, and the sheet is used for lookup.
"""

from __future__ import annotations

import argparse
import re
import sys
from collections import OrderedDict
from dataclasses import dataclass, field
from pathlib import Path

import docx
from docx.opc.exceptions import PackageNotFoundError

import docx_blocks as db
import export_vocab_docx as ev
from export_docx_splice import assert_student_edition_is_clean

MERGED = "重难点词汇表_合并_学生版.docx"


def _key(header_cell: str) -> str:
    """Headers are matched with the spacing removed — including the non-breaking kind.

    A handout that was opened and re-saved in Word can come back with a stray space or an
    nbsp in a header cell, and refusing to recognise it would silently drop that whole
    paper's words.
    """
    return re.sub(r"\s", "", header_cell.replace("\xa0", ""))


_WORDS_HEADER = tuple(_key(h) for h in ev.WORDS_HEADER)
_FORMS_HEADER = tuple(_key(h) for h in ev.FORMS_HEADER)

_PAREN = re.compile(r"[（(][^）)]*[）)]")
_SENSE_SPLIT = "；;，,、"


@dataclass
class Source:
    """One input file, and what was found in it."""

    path: Path
    reading: list[dict] = field(default_factory=list)
    forms: list[dict] = field(default_factory=list)
    skipped: str = ""  # why nothing was taken from it, if nothing was

    @property
    def used(self) -> bool:
        return bool(self.reading or self.forms)


@dataclass
class Report:
    sources: list[Source]
    reading: list[dict]
    forms: list[dict]
    out: Path | None = None

    @property
    def reading_before(self) -> int:
        return sum(len(s.reading) for s in self.sources)

    @property
    def forms_before(self) -> int:
        return sum(len(s.forms) for s in self.sources)

    def lines(self) -> list[str]:
        out = []
        for source in self.sources:
            if source.used:
                out.append(f"  {source.path.name}：阅读词汇 {len(source.reading)}，语法变形 {len(source.forms)}")
            else:
                out.append(f"  {source.path.name}：跳过（{source.skipped}）")
        out.append(
            f"阅读词汇 {self.reading_before} → {len(self.reading)}"
            f"（去掉重复 {self.reading_before - len(self.reading)}）"
        )
        out.append(
            f"语法变形 {self.forms_before} → {len(self.forms)}"
            f"（去掉重复 {self.forms_before - len(self.forms)}）"
        )
        if self.out:
            out.append(f"已写出：{self.out}")
        return out


# --- reading the handouts


def _rows(table) -> list[list[str]]:
    """Every row's cell text, walking the XML directly.

    `table.rows[i].cells` rebuilds the whole cell grid per row, which is quadratic — the
    same trap `export_vocab_docx._add_table` used to fall into, and these tables are the
    long ones. Neither handout table has merged cells, so a row's `tc` children *are*
    its cells.

    The text comes from `docx_blocks.node_text` (which collects `w:t`), not lxml's
    `itertext()`: on python-docx's element classes the latter yields each run's text
    three times over, so 「英文单词」 arrives as 「英文单词英文单词英文单词」 and no header
    is ever recognised.
    """
    return [[db.node_text(tc).strip() for tc in tr.tc_lst] for tr in table._tbl.tr_lst]


def read_handout(path: Path) -> Source:
    """Pull the two known tables out of one handout, by header.

    Unrecognised tables are skipped rather than guessed at: a 3-column table is not
    necessarily a word list, and a wrong row here would put a word on a student's sheet
    that no model ever chose. The caller is told what was skipped.
    """
    source = Source(path=path)
    try:
        document = docx.Document(str(path))
    except PackageNotFoundError:
        # python-docx says "Package not found at '<full path>'", which in the UI is a line
        # of temp path where the teacher needs one word. The file name is already shown.
        source.skipped = "不是有效的 Word 文件"
        return source
    except Exception as exc:  # noqa: BLE001 - python-docx raises several unrelated types
        source.skipped = f"打不开：{str(exc).splitlines()[0][:80]}"
        return source

    unknown = 0
    for table in document.tables:
        rows = _rows(table)
        if not rows:
            continue
        header = tuple(_key(cell) for cell in rows[0])
        if header == _WORDS_HEADER:
            for cells in rows[1:]:
                if any(cells) and cells[0].strip():
                    source.reading.append(
                        {"word": cells[0], "pos": cells[1], "meaning": cells[2]}
                    )
        elif header == _FORMS_HEADER:
            for cells in rows[1:]:
                if any(cells) and cells[0].strip() and cells[2].strip():
                    source.forms.append({
                        "base": cells[0], "base_pos": cells[1],
                        "derived": cells[2], "derived_pos": cells[3], "note": cells[4],
                    })
        else:
            unknown += 1

    if not source.used:
        source.skipped = (
            f"没找到词汇表（{unknown} 张表的表头都不认识）" if unknown else "里面没有表格"
        )
    return source


def handouts_in(paths: list[Path], exclude: Path | None = None) -> list[Path]:
    """Expand what was dropped on us into a list of .docx to read.

    Folders are searched one level deep, Word's own ``~$`` lock files are ignored, and the
    file we are about to write is never read: merging a previous merge back into itself
    would keep working, but the counts would stop meaning anything.
    """
    found: list[Path] = []
    for path in paths:
        candidates = sorted(path.glob("*.docx")) if path.is_dir() else [path]
        for candidate in candidates:
            if candidate.name.startswith("~$") or candidate.suffix.lower() != ".docx":
                continue
            if exclude and candidate.resolve() == exclude.resolve():
                continue
            if candidate not in found:
                found.append(candidate)
    return found


# --- dedupe


def merge_pos(values: list[str]) -> str:
    """Union of 词性 tags: 'v./n.' + 'n.' -> 'v./n.'; 'n.' + 'adj.' -> 'n./adj.'."""
    toks: list[str] = []
    for value in values:
        for tok in re.split(r"[/、]|\s*&\s*", value):
            tok = tok.strip()
            if tok and tok not in toks:
                toks.append(tok)
    # 'n. phr.' already covers 'n.' and 'phr.' — keep the specific one, drop what it spans.
    keep = [t for t in toks
            if not any(other != t and (other.startswith(t) or other.endswith(t)) for other in toks)]
    return "/".join(keep or toks)


def _sense_key(sense: str) -> str:
    """What makes two glosses the same sense: the words, without notes or punctuation."""
    return re.sub(r"[\s.,，、；;：:\-—“”\"']", "", _PAREN.sub("", sense)).lower()


def split_senses(meaning: str) -> list[str]:
    """Split a gloss into senses without cutting inside （…） — the notes hold commas."""
    out, depth, buf = [], 0, ""
    for char in meaning:
        if char in "（(":
            depth += 1
        elif char in "）)":
            depth = max(0, depth - 1)
        if depth == 0 and char in _SENSE_SPLIT:
            out.append(buf)
            buf = ""
            continue
        buf += char
    out.append(buf)
    return [sense.strip() for sense in out if sense.strip()]


def merge_meaning(values: list[str]) -> str:
    """Union the senses across papers; the annotated wording of a sense wins."""
    senses: OrderedDict[str, str] = OrderedDict()
    for value in values:
        for sense in split_senses(value):
            key = _sense_key(sense)
            if not key:
                continue
            if key not in senses or len(sense) > len(senses[key]):
                senses[key] = sense
    return "；".join(senses.values())


def merge_note(values: list[str]) -> str:
    """Keep the most informative single 考点说明 — see the module docstring."""
    variants = sorted({value.strip() for value in values if value.strip()}, key=lambda v: (-len(v), v))
    return variants[0] if variants else ""


def merge_entries(sources: list[Source]) -> tuple[list[dict], list[dict]]:
    by_word: OrderedDict[str, list[dict]] = OrderedDict()
    by_form: OrderedDict[tuple[str, str], list[dict]] = OrderedDict()
    for source in sources:
        for entry in source.reading:
            by_word.setdefault(entry["word"].strip().lower(), []).append(entry)
        for entry in source.forms:
            key = (entry["base"].strip().lower(), entry["derived"].strip().lower())
            by_form.setdefault(key, []).append(entry)

    reading = [{
        "word": group[0]["word"].strip(),
        "pos": merge_pos([e["pos"].strip() for e in group]),
        "meaning": merge_meaning([e["meaning"].strip() for e in group]),
    } for group in (by_word[key] for key in sorted(by_word))]

    forms = [{
        "base": group[0]["base"].strip(),
        "base_pos": merge_pos([e["base_pos"].strip() for e in group]),
        "derived": group[0]["derived"].strip(),
        "derived_pos": merge_pos([e["derived_pos"].strip() for e in group]),
        "note": merge_note([e["note"].strip() for e in group]),
    } for group in (by_form[key] for key in sorted(by_form))]

    return reading, forms


# --- the whole job


def merge_handouts(paths: list[Path], out: Path, log=print) -> Report:
    """Read every handout in `paths`, write one deduped handout to `out`."""
    files = handouts_in(paths, exclude=out)
    if not files:
        raise RuntimeError("没有找到 .docx 词汇表文件")

    sources = [read_handout(path) for path in files]
    reading, forms = merge_entries(sources)
    if not reading and not forms:
        raise RuntimeError(
            "这些文件里没有认得出来的词汇表。"
            "需要的是本工具生成的「重难点词汇表」（表头为「英文单词/词性/准确的中文释义」"
            "或「基础词汇/…/考点说明」）。"
        )

    report = Report(sources=sources, reading=reading, forms=forms)
    for line in report.lines():
        log(line)

    ev.build([{"reading_words": reading, "word_forms": forms}], out)
    # The merged sheet is a student handout, so it goes through the same last gate the
    # exported one does — a teacher edition dropped in by mistake must not ride along.
    assert_student_edition_is_clean(out)
    report.out = out
    log(f"已写出：{out}")
    return report


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="把多份已生成的重难点词汇表合并成一份，去掉重复（纯本地，不调用 AI）"
    )
    parser.add_argument("inputs", nargs="+", type=Path, help="词汇表 .docx 文件，或装着它们的文件夹")
    parser.add_argument("--out", type=Path, help=f"输出文件（默认：第一个输入所在目录下的 {MERGED}）")
    args = parser.parse_args(argv)

    first = args.inputs[0]
    out = args.out or (first if first.is_dir() else first.parent) / MERGED
    try:
        merge_handouts(args.inputs, out)
    except RuntimeError as exc:
        print(f"失败：{exc}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
