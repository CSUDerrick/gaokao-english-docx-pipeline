#!/usr/bin/env python3
"""Clone-and-splice docx writer.

The old export rebuilt a Word file from Markdown, which meant every font, indent,
bold run and image in the teacher's original paper had to be re-invented — and
images simply vanished.  This module instead *copies the original paragraphs*.

The one hard problem is that selected questions come from several different exam
papers, and Chinese Word files all use auto-generated style ids that collide:
``a3`` means ``footer`` in one paper, ``Title`` in another, and ``rId6`` is an
image in one and ``endnotes.xml`` in another.  Naively moving paragraphs between
documents silently renders a footer as a headline and points pictures at notes.

So the work is split in two, and neither half has to solve that:

1. :func:`clone_subset` — copy ONE source .docx byte-for-byte and keep only the
   chosen paragraphs.  Single source, so every style id and relationship id is
   still the one it was authored against.  Images, tables and fonts come along
   for free because ``word/media/*``, ``styles.xml`` and ``_rels`` are untouched.
2. :func:`merge` — concatenate those single-source files with ``docxcompose``,
   which exists to remap exactly these id collisions.

New content (the AI explanations) has no original to clone, so it is written
with python-docx against named styles that are *registered in styles.xml before
being referenced* — see ``docs/word_compatibility.md`` for why that matters.
"""

from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Word chokes on these and reports "unreadable content" — see word_compatibility.md
ILLEGAL_XML = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")

# Word prints text verbatim, so any Markdown that reaches it is just noise on the
# page ("**原句：**"). Underscores are deliberately left alone: "____36____" is a
# fill-in-the-blank, not emphasis — see test_notes_strip_illegal_control_characters.
MARKDOWN_NOISE = re.compile(r"\*\*|`|^\s{0,3}#{1,6}\s+", re.MULTILINE)

NOTE_BODY = "讲解正文"
NOTE_HEADING = "讲解标题"
NOTE_ANSWER = "讲解答案"

# The printed vocabulary handout has its own typography (Times New Roman + 宋体,
# 三号 title, 小四 body) and must not drag the teacher edition's notes along with it.
VOCAB_TITLE = "词汇标题"
VOCAB_BODY = "词汇正文"


def sanitize(text: str) -> str:
    return MARKDOWN_NOISE.sub("", ILLEGAL_XML.sub("", text or ""))


# CT_PPrBase / CT_SectPr are sequences too. The Pandoc-derived templates emit
# w:keepNext and w:pageBreakBefore *after* the properties they must precede, and
# their w:pgMar has no w:gutter, which the schema requires.
P_PR_ORDER = (
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr", "widowControl",
    "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs", "suppressAutoHyphens",
    "kinsoku", "wordWrap", "overflowPunct", "topLinePunct", "autoSpaceDE",
    "autoSpaceDN", "bidi", "adjustRightInd", "snapToGrid", "spacing", "ind",
    "contextualSpacing", "mirrorIndents", "suppressOverlap", "jc", "textDirection",
    "textAlignment", "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr",
    "sectPr", "pPrChange",
)


def _reorder(pr, order: tuple[str, ...]) -> None:
    """Sort a properties element's children into the schema's required sequence."""
    def rank(child) -> int:
        return order.index(child.tag.rsplit("}", 1)[-1])

    known = [c for c in pr if isinstance(c.tag, str) and c.tag.rsplit("}", 1)[-1] in order]
    if [rank(c) for c in known] == sorted(rank(c) for c in known):
        return
    for child in known:
        pr.remove(child)
    for child in sorted(known, key=rank):
        pr.append(child)


def normalize_template_ooxml(doc) -> None:
    """Repair the schema violations the reference templates were built with."""
    for part in (doc.element, doc.styles.element):
        for pr in part.iter(f"{{{W}}}pPr"):
            _reorder(pr, P_PR_ORDER)
        for margin in part.iter(f"{{{W}}}pgMar"):
            if margin.get(f"{{{W}}}gutter") is None:
                margin.set(f"{{{W}}}gutter", "0")


def blank_template(template: Path):
    """Open a reference template and strip its sample content, keeping page setup.

    The Pandoc-derived templates ship with demo paragraphs *and a demo table*.
    Deleting only ``doc.paragraphs`` left the table behind — a stray "Table 1 2"
    grid rode along in every answer sheet. Page size is forced rather than
    inherited so a missing template costs fonts, never A4.
    """
    from docx import Document
    from docx.shared import Cm

    doc = Document(str(template)) if template.exists() else Document()
    body = doc.element.body
    for child in list(body):
        if isinstance(child.tag, str) and not child.tag.endswith("}sectPr"):
            body.remove(child)
    for section in doc.sections:
        section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    normalize_template_ooxml(doc)
    return doc


# CT_TblPrBase is a *sequence*: its children have to appear in this order, or Word
# rejects the file. python-docx's add_table() emits tblW + tblLook, so appending
# tblBorders at the end puts it after tblLook — which is invalid.
TBL_PR_ORDER = (
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
    "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd", "tblBorders",
    "shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption", "tblDescription",
)


def set_table_borders(table) -> None:
    """Draw a table's gridlines directly, without naming a style.

    Deliberately not ``table.style = "Table Grid"``: that style is not guaranteed
    to exist in every template, and referencing a style styles.xml never declares
    is one of the ways Word ends up reporting unreadable content.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        line = OxmlElement(f"w:{edge}")
        line.set(qn("w:val"), "single")
        line.set(qn("w:sz"), "4")
        line.set(qn("w:color"), "000000")
        borders.append(line)

    tbl_pr = table._tbl.tblPr
    rank = TBL_PR_ORDER.index("tblBorders")
    for existing in tbl_pr:
        name = existing.tag.rsplit("}", 1)[-1]
        if name in TBL_PR_ORDER and TBL_PR_ORDER.index(name) > rank:
            existing.addprevious(borders)
            return
    tbl_pr.append(borders)


def _element_children(body) -> list:
    return [c for c in body if isinstance(c.tag, str)]


# Range markers come in pairs that may straddle paragraphs. Subsetting a body by
# paragraph index can keep one end and drop the other, and Word reports the
# surviving half as "unreadable content" (this is what OLE_LINK7/OLE_LINK8 did to
# the teacher edition: 8 bookmarkStart, 6 bookmarkEnd).
RANGE_MARKERS = [
    ("bookmarkStart", "bookmarkEnd", f"{{{W}}}id"),
    ("commentRangeStart", "commentRangeEnd", f"{{{W}}}id"),
    ("moveFromRangeStart", "moveFromRangeEnd", f"{{{W}}}id"),
    ("moveToRangeStart", "moveToRangeEnd", f"{{{W}}}id"),
    ("permStart", "permEnd", f"{{{W}}}id"),
]


def drop_orphan_range_markers(body) -> int:
    """Remove range markers whose partner is not in ``body``. Returns the count."""
    removed = 0
    for start_tag, end_tag, id_attr in RANGE_MARKERS:
        starts = list(body.iter(f"{{{W}}}{start_tag}"))
        ends = list(body.iter(f"{{{W}}}{end_tag}"))
        start_ids = {el.get(id_attr) for el in starts}
        end_ids = {el.get(id_attr) for el in ends}
        for el in starts:
            if el.get(id_attr) not in end_ids:
                el.getparent().remove(el)
                removed += 1
        for el in ends:
            if el.get(id_attr) not in start_ids:
                el.getparent().remove(el)
                removed += 1
    return removed


def orphan_range_markers(doc) -> list[str]:
    """Name every range marker in ``doc`` (a parsed document.xml) missing its partner."""
    orphans: list[str] = []
    for start_tag, end_tag, id_attr in RANGE_MARKERS:
        start_ids = {el.get(id_attr) for el in doc.iter(f"{{{W}}}{start_tag}")}
        end_ids = {el.get(id_attr) for el in doc.iter(f"{{{W}}}{end_tag}")}
        orphans += [f"{start_tag}[@id={i}]" for i in sorted(start_ids - end_ids) if i is not None]
        orphans += [f"{end_tag}[@id={i}]" for i in sorted(end_ids - start_ids) if i is not None]
    return orphans


def clone_subset(src: Path, body_indices: list[int], out: Path) -> Path:
    """Copy ``src`` keeping only the body children in ``body_indices``.

    Everything except ``word/document.xml`` is copied verbatim, so styles, theme,
    fonts, numbering, relationships and ``word/media/*`` all still resolve.  The
    original ``w:sectPr`` is kept so page size and margins survive.
    """
    wanted = sorted(set(body_indices))

    with zipfile.ZipFile(src) as zin:
        doc_xml = zin.read("word/document.xml")
        entries = [(i, zin.read(i.filename)) for i in zin.infolist()]

    root = etree.fromstring(doc_xml)
    body = root.find(f"{{{W}}}body")
    if body is None:
        raise RuntimeError(f"{src} has no w:body")

    children = _element_children(body)
    sect_pr = body.find(f"{{{W}}}sectPr")

    keep = [children[i] for i in wanted if 0 <= i < len(children)]
    # sectPr is a body child too; it must stay last and must not be duplicated.
    keep = [c for c in keep if c is not sect_pr]

    for child in list(body):
        body.remove(child)
    for child in keep:
        body.append(child)
    if sect_pr is not None:
        body.append(sect_pr)

    drop_orphan_range_markers(body)

    new_doc = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

    out.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for info, data in entries:
            zout.writestr(info, new_doc if info.filename == "word/document.xml" else data)
    return out


def ensure_note_styles(doc: Document) -> None:
    """Register the explanation styles before anything references them.

    Named styles (rather than direct run formatting) so the teacher can restyle
    every explanation in one go via Word's style pane.
    """
    specs = [
        (NOTE_HEADING, 12, True),
        (NOTE_BODY, 10.5, False),
        (NOTE_ANSWER, 10.5, True),
    ]
    existing = {s.name for s in doc.styles}
    for name, size, bold in specs:
        if name in existing:
            continue
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        # East-Asian font must be set on rPr/rFonts or Word falls back for Chinese.
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")


def _paragraph_layout(pr, *, centred: bool = False, page_break: bool = False) -> None:
    """The paragraph settings the printed handout needs, written straight into pPr.

    Three of these are Word checkboxes the teacher asked to be off, and each is an
    *absence* rather than a value — so they can only be got right by writing pPr
    ourselves:

    * 「定义文档网格时对齐网格」 → ``w:snapToGrid val=0``. On, Word stretches the
      line height to the document grid and the tight table rows drift apart.
    * 「在相同样式的段落间不添加空格」 → ``w:contextualSpacing``, which we simply
      never emit.
    * 段前/段后 0 行 → ``w:spacing before=0 after=0``, single-spaced lines.

    pPr is a *sequence* (决策 11), so everything is appended and then sorted into
    P_PR_ORDER rather than trusted to land in the right place.
    """
    def child(tag: str) -> object:
        found = pr.find(qn(f"w:{tag}"))
        if found is None:
            found = OxmlElement(f"w:{tag}")
            pr.append(found)
        return found

    if page_break:
        child("pageBreakBefore")

    child("snapToGrid").set(qn("w:val"), "0")

    spacing = child("spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")

    if centred:
        child("jc").set(qn("w:val"), "center")

    _reorder(pr, P_PR_ORDER)


def ensure_vocab_styles(doc: Document) -> None:
    """The vocabulary handout's own styles: Times New Roman + 宋体, printable.

    Deliberately *not* the NOTE_* styles — those set the teacher edition's notes in
    Arial Unicode MS at 10.5pt, and restyling them to suit a printed word list would
    silently restyle every explanation in the teacher edition too.
    """
    specs = [
        (VOCAB_TITLE, 16, True, True),   # 三号, 加粗, 居中
        (VOCAB_BODY, 12, False, False),  # 小四
    ]
    existing = {s.name for s in doc.styles}
    for name, size, bold, centred in specs:
        if name in existing:
            continue
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        fonts = style.element.rPr.rFonts
        fonts.set(qn("w:eastAsia"), "宋体")
        fonts.set(qn("w:cs"), "Times New Roman")
        # python-docx writes w:sz but not w:szCs; the reference handout carries both.
        size_cs = OxmlElement("w:szCs")
        size_cs.set(qn("w:val"), str(int(size * 2)))
        style.element.rPr.append(size_cs)
        _paragraph_layout(style.element.get_or_add_pPr(), centred=centred)


def fit_table_to_window(table) -> None:
    """Make the table span the text column and share it evenly, whatever the content.

    ``w:tblW type=pct w:w=5000`` is Word's 「根据窗口自动调整」. Without it python-docx
    writes fixed twip widths measured off the *template's* page, so a table built on
    one page size prints half-width on another.

    The table style is dropped at the same time: the borders are drawn directly
    (see ``set_table_borders``), and a ``w:tblStyle`` naming a style the document
    never declares is one of the ways Word ends up reporting unreadable content.
    """
    tbl_pr = table._tbl.tblPr
    for style_ref in tbl_pr.findall(qn("w:tblStyle")):
        tbl_pr.remove(style_ref)

    width = tbl_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), "5000")
    width.set(qn("w:type"), "pct")

    columns = len(table.columns)
    share = str(5000 // columns)
    for row in table.rows:
        for cell in row.cells:
            cell_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(cell_width)
            cell_width.set(qn("w:w"), share)
            cell_width.set(qn("w:type"), "pct")

    _reorder(tbl_pr, TBL_PR_ORDER)


def drop_headers(doc: Document) -> None:
    """Remove the header the reference template carries.

    ``blank_template`` keeps the template's sectPr so page setup survives — and the
    sectPr carries a ``w:headerReference``, which is how the student template's
    「高三英语精选训练」 header ended up on top of the vocabulary handout. The teacher
    prints this one, and asked for no header.
    """
    for sect_pr in doc.element.body.iter(qn("w:sectPr")):
        for ref in sect_pr.findall(qn("w:headerReference")):
            sect_pr.remove(ref)


def set_page_margins(doc: Document, twips: int) -> None:
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Twips(twips)
        section.left_margin = section.right_margin = Twips(twips)


def page_break_before(paragraph) -> None:
    """Start this paragraph on a new page.

    The teacher's own handout does this with seventeen empty paragraphs, which only
    lands on the right page for the exact word count it was built with. A real
    pageBreakBefore survives the list getting longer or shorter.
    """
    _paragraph_layout(paragraph._p.get_or_add_pPr(), page_break=True)


def decorate(path: Path, heading: str = "", notes: list[tuple[str, str]] | None = None) -> None:
    """Add a heading above, and explanation paragraphs below, a cloned question."""
    notes = notes or []
    if not heading and not notes:
        return
    doc = Document(str(path))
    ensure_note_styles(doc)
    if heading:
        first = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        first.insert_paragraph_before(sanitize(heading), style=NOTE_HEADING)
    for style_name, text in notes:
        for line in sanitize(text).split("\n"):
            doc.add_paragraph(line, style=style_name)
    doc.save(str(path))


AUTHOR = "高三英语试卷整理工具"


def scrub_metadata(path: Path, title: str) -> None:
    """Drop the original paper's author out of the exported file.

    The clone copies ``docProps`` verbatim, which carries the name of whoever
    authored the source exam — a real person — into the teacher's output.

    Two things make this fiddlier than it looks, and both have bitten:

    * python-docx's ``core_properties`` setter does not reach the core.xml that
      docxcompose actually writes, so the name survived a merge.
    * core.xml can end up holding a *second* ``lastModifiedBy`` whose ``cp:``
      prefix is bound to the custom-properties namespace rather than
      core-properties, so a namespace-qualified lookup misses it.

    So: rewrite the part in the zip, and match on local name only.

    A field is *removed*, never blanked. ``cp:lastPrinted`` is typed as a
    dateTime, so emptying it leaves ``<cp:lastPrinted></cp:lastPrinted>`` — an
    empty string where a date must be. That is invalid, and Word answers with
    "发现无法读取的内容". It is also why only the student and teacher editions
    complained: the answer sheet is built from a template that has no
    lastPrinted to blank in the first place.
    """
    replace = {"creator": AUTHOR, "lastModifiedBy": AUTHOR, "title": title}
    remove = {"description", "lastPrinted", "category", "keywords"}

    with zipfile.ZipFile(path) as zin:
        entries = [(i, zin.read(i.filename)) for i in zin.infolist()]

    out: list[tuple[zipfile.ZipInfo, bytes]] = []
    for info, data in entries:
        if info.filename == "docProps/core.xml":
            root = etree.fromstring(data)
            for el in list(root.iter()):
                if not isinstance(el.tag, str):
                    continue
                name = el.tag.rsplit("}", 1)[-1]
                if name in replace:
                    el.text = replace[name]
                elif name in remove and el.getparent() is not None:
                    el.getparent().remove(el)
            data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
        out.append((info, data))

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for info, data in out:
            zout.writestr(info, data)


def merge(parts: list[Path], out: Path, page_break: bool = True) -> Path:
    """Concatenate single-source docx files, remapping colliding ids."""
    if not parts:
        raise ValueError("nothing to merge")

    from docxcompose.composer import Composer

    master = Document(str(parts[0]))
    composer = Composer(master)
    for part in parts[1:]:
        composer.append(Document(str(part)), remove_property_fields=True)

    out.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(out))
    return out


def media_count(path: Path) -> int:
    with zipfile.ZipFile(path) as z:
        return len([n for n in z.namelist() if n.startswith("word/media/")])


def validate(path: Path) -> None:
    """Fail loudly rather than hand the teacher a file Word wants to repair."""
    with zipfile.ZipFile(path) as z:
        for name in z.namelist():
            if name.endswith((".xml", ".rels")):
                etree.fromstring(z.read(name))  # raises on malformed XML
        doc = etree.fromstring(z.read("word/document.xml"))
        styles = etree.fromstring(z.read("word/styles.xml"))

    pg = doc.find(f".//{{{W}}}sectPr/{{{W}}}pgSz")
    if pg is None:
        raise RuntimeError(f"{path.name}: no page size")
    width, height = int(pg.get(f"{{{W}}}w")), int(pg.get(f"{{{W}}}h"))
    if not (11800 <= width <= 12050 and 16700 <= height <= 17000):
        raise RuntimeError(f"{path.name}: page is {width}x{height} twips, expected A4 portrait")

    defined = {s.get(f"{{{W}}}styleId") for s in styles.findall(f"{{{W}}}style")}
    used = {
        e.get(f"{{{W}}}val")
        for tag in ("pStyle", "rStyle", "tblStyle")
        for e in doc.findall(f".//{{{W}}}{tag}")
    }
    if not used <= defined:
        raise RuntimeError(f"{path.name}: references undefined styles {used - defined}")

    dangling = unresolved_rids(path)
    if dangling:
        raise RuntimeError(f"{path.name}: unresolved relationship ids {dangling}")

    orphans = orphan_range_markers(doc)
    if orphans:
        raise RuntimeError(f"{path.name}: unpaired range markers {orphans}")

    empty = empty_typed_core_properties(path)
    if empty:
        raise RuntimeError(f"{path.name}: docProps/core.xml has empty date fields {empty}")


# cp:lastPrinted and the dcterms fields are dateTimes. An empty one is not "no
# date", it is a malformed date, and Word refuses the file over it.
DATE_CORE_PROPERTIES = ("lastPrinted", "created", "modified")


def empty_typed_core_properties(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as z:
        if "docProps/core.xml" not in z.namelist():
            return []
        root = etree.fromstring(z.read("docProps/core.xml"))
    return [
        el.tag.rsplit("}", 1)[-1]
        for el in root.iter()
        if isinstance(el.tag, str)
        and el.tag.rsplit("}", 1)[-1] in DATE_CORE_PROPERTIES
        and not (el.text or "").strip()
    ]


def unresolved_rids(path: Path) -> list[str]:
    """Relationship ids referenced by the body but absent from the rels part.

    A non-empty result means a picture points at nothing — the exact corruption
    that cross-document splicing causes if ids are not remapped.
    """
    with zipfile.ZipFile(path) as z:
        doc = z.read("word/document.xml").decode("utf-8")
        try:
            rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        except KeyError:
            rels = ""
    used = set(re.findall(r'r:(?:embed|id|link)="([^"]+)"', doc))
    defined = set(re.findall(r'Id="([^"]+)"', rels))
    return sorted(used - defined)
