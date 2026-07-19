#!/usr/bin/env python3
"""PDF input via PaddleOCR-VL.

A scanned paper has no OOXML to clone, so it has to be typeset. Rather than add a
second code path through the whole pipeline, OCR output is written out as a real
``.docx`` first — after that it is just another source paper: it gets segmented,
scored, cloned and exported by exactly the same code as a Word input.

    paper.pdf --[PaddleOCR-VL]--> paper.docx --> (the normal pipeline)

The generated .docx is what ``source_blocks`` then points at, so the "clone the
original" export works for scanned papers too; the layout it clones is the one
reconstructed here.

Endpoint: one global **async job** API — submit the file, poll the job, download the
result JSONL. That is the same submit/poll/download shape ``mineru_ingest`` already
uses, and the reason this module looks like it. The token comes from
https://aistudio.baidu.com/account/accessToken and is read from
``PADDLEOCR_ACCESS_TOKEN``; ``PADDLEOCR_BASE_URL`` is now only an override for a
private deployment — an old per-account ``.../layout-parsing`` URL left in the
Keychain still works, because only its host is kept.
"""

from __future__ import annotations

import base64
import json
import os
import time
import urllib.error
import urllib.parse
import urllib.request
import uuid
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.shared import Inches

import docx_splice as ds
from bundle_paths import template_dir

MODEL = "PaddleOCR-VL-1.6"
TOKEN_ENV = "PADDLEOCR_ACCESS_TOKEN"
BASE_URL_ENV = "PADDLEOCR_BASE_URL"

# v2 is one global service, so the teacher no longer has to hunt for a per-account URL
# in the console — the old code demanded one and could not run without it.
DEFAULT_BASE = "https://paddleocr.aistudio-app.com"
JOBS_PATH = "/api/v2/ocr/jobs"
POLL_INTERVAL = 5.0
POLL_TIMEOUT = 900.0

# PaddleOCR block_label -> what it should look like in Word
HEADINGS = {"doc_title", "paragraph_title", "title"}
SKIP = {"header", "footer", "page_number", "seal"}


@dataclass
class OcrBlock:
    label: str
    text: str
    order: int
    image: bytes | None = None
    bbox: tuple[float, float, float, float] | None = None


@dataclass
class OcrResult:
    blocks: list[OcrBlock] = field(default_factory=list)
    markdown: str = ""


def _jobs_url(base_url: str) -> str:
    """The job endpoint. An empty base URL is the normal case now.

    v2 is one global service, so a teacher who configured nothing still works. A base
    URL that *is* set is treated as a host override — including an old
    ``.../layout-parsing`` URL still sitting in the Keychain, whose path is dropped
    rather than pasted onto the new one.
    """
    base = (base_url or "").strip()
    if not base:
        return f"{DEFAULT_BASE}{JOBS_PATH}"
    if JOBS_PATH in base:
        return base.split(JOBS_PATH)[0].rstrip("/") + JOBS_PATH
    parts = urllib.parse.urlsplit(base if "//" in base else f"https://{base}")
    return f"{parts.scheme or 'https'}://{parts.netloc}{JOBS_PATH}"


def _encode_multipart(fields: dict[str, str], name: str, filename: str, blob: bytes) -> tuple[str, bytes]:
    """Build a multipart/form-data body by hand.

    The API's own examples use requests' ``files=``. This project has no requests — no
    HTTP library at all beyond the stdlib — and is not about to grow one just to upload
    a file: urllib plus twenty lines does the same job and keeps the frozen .app the
    size it is.
    """
    boundary = "----" + uuid.uuid4().hex
    body = bytearray()
    for key, value in fields.items():
        body += f"--{boundary}\r\n".encode()
        body += f'Content-Disposition: form-data; name="{key}"\r\n\r\n'.encode()
        body += f"{value}\r\n".encode()
    body += f"--{boundary}\r\n".encode()
    body += f'Content-Disposition: form-data; name="{name}"; filename="{filename}"\r\n'.encode()
    body += b"Content-Type: application/pdf\r\n\r\n"
    body += blob + b"\r\n"
    body += f"--{boundary}--\r\n".encode()
    return f"multipart/form-data; boundary={boundary}", bytes(body)


def _request(url: str, token: str, *, data: bytes | None = None,
             content_type: str = "", timeout: int = 120) -> dict:
    """Every call to the service goes through here — the one seam the tests replace.

    Auth is ``bearer``. The old per-account service wanted ``token <t>``; sending that
    to v2 is a 401, which is half of why every OCR run was failing.
    """
    headers = {"Authorization": f"bearer {token}"}
    if content_type:
        headers["Content-Type"] = content_type
    request = urllib.request.Request(
        url, data=data, headers=headers, method="POST" if data else "GET"
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            body = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", "replace")[:400]
        if exc.code in (401, 403):
            raise RuntimeError(
                f"PaddleOCR 令牌无效（HTTP {exc.code}）。"
                "请到 https://aistudio.baidu.com/account/accessToken 重新获取。"
            ) from exc
        if exc.code == 429:
            raise RuntimeError("PaddleOCR 免费额度已用完（HTTP 429），请稍后再试或在控制台查看配额。") from exc
        raise RuntimeError(f"PaddleOCR API 调用失败 (HTTP {exc.code}): {detail}") from exc

    if body.get("code") not in (0, None):
        raise RuntimeError(f"PaddleOCR 返回错误：{body.get('msg') or body.get('code')}")
    return body


def _submit(pdf: Path, token: str, jobs_url: str) -> str:
    optional = {
        "useDocOrientationClassify": False,
        "useDocUnwarping": False,
        # Exam papers carry charts often enough to be worth reading.
        "useChartRecognition": True,
    }
    content_type, body = _encode_multipart(
        {"model": MODEL, "optionalPayload": json.dumps(optional)},
        "file", pdf.name, pdf.read_bytes(),
    )
    reply = _request(jobs_url, token, data=body, content_type=content_type, timeout=300)
    job_id = str((reply.get("data") or {}).get("jobId") or "")
    if not job_id:
        raise RuntimeError("PaddleOCR 收下了文件却没有给 jobId。")
    return job_id


def _wait(job_id: str, token: str, jobs_url: str, *, sleep, log=None) -> str:
    """Poll until the job is done, and return the result JSONL url.

    ``sleep`` is injected rather than called directly: that is the only thing that makes
    取消 work. A worker parked in ``time.sleep`` would keep the run alive until the OCR
    finished no matter what the teacher pressed. mineru_ingest waits the same way, for
    the same reason.
    """
    deadline = time.time() + POLL_TIMEOUT
    while True:
        data = _request(f"{jobs_url}/{job_id}", token).get("data") or {}
        state = str(data.get("state") or "")
        if state == "done":
            url = str((data.get("resultUrl") or {}).get("jsonUrl") or "")
            if not url:
                raise RuntimeError("PaddleOCR 说解析完成了，却没有给结果地址。")
            return url
        if state == "failed":
            raise RuntimeError(f"PaddleOCR 解析失败：{data.get('errorMsg') or '未说明原因'}")
        if log:
            progress = data.get("extractProgress") or {}
            if progress.get("totalPages"):
                log(f"    OCR 中：{progress.get('extractedPages')}/{progress.get('totalPages')} 页")
        if time.time() > deadline:
            raise RuntimeError(f"PaddleOCR 解析超时（超过 {POLL_TIMEOUT / 60:.0f} 分钟）。")
        sleep(POLL_INTERVAL)


def _download_jsonl(url: str, timeout: int = 300) -> list[str]:
    with urllib.request.urlopen(url, timeout=timeout) as response:
        text = response.read().decode("utf-8")
    return [line for line in text.strip().split("\n") if line.strip()]


def _image_bytes(value: object) -> bytes | None:
    """A figure arrives either as a URL to fetch or as base64 — accept both.

    The old synchronous API inlined images as base64; v2 hands back a URL. Guessing
    wrong just loses the figure, so try the cheap discriminator and keep going: an
    undecodable figure must not sink the paper.
    """
    if not isinstance(value, str) or not value:
        return None
    try:
        if value.startswith(("http://", "https://")):
            with urllib.request.urlopen(value, timeout=120) as response:
                return response.read()
        return base64.b64decode(value)
    except Exception:
        return None


def check_service(base_url: str, token: str, timeout: int = 20) -> tuple[bool, str]:
    """Is this token usable? Answered without uploading anything or spending quota.

    Asking for a job id that cannot exist is enough to tell the two failures apart: a
    good token gets "jobId 不存在" (HTTP 404, code 11001) — the service authenticated us
    and then went looking — while a bad token never gets that far and is refused 401.
    """
    request = urllib.request.Request(
        f"{_jobs_url(base_url)}/1",
        headers={"Authorization": f"bearer {token}"},
        method="GET",
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            response.read(200)
        return True, "服务可访问"
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", "replace")[:200]
        if exc.code in (401, 403):
            return False, f"令牌无效（HTTP {exc.code}）"
        if exc.code == 429:
            return False, "免费额度已用完（HTTP 429）"
        if exc.code == 404:
            # It looked the job up, so both the address and the token are fine.
            if "jobId" in detail or "11001" in detail:
                return True, "服务可访问，令牌有效"
            return False, "服务地址不对（HTTP 404）"
        return True, f"服务可访问（HTTP {exc.code}）"
    except urllib.error.URLError as exc:
        return False, f"连不上服务：{exc.reason}"


def parse_response(body: dict, page_offset: int = 0) -> OcrResult:
    """Flatten one result object into reading-ordered blocks.

    The response nests one entry per page; ``parsing_res_list`` inside each holds
    the layout blocks with their reading order, which is what decides the order
    paragraphs land in the Word file. ``page_offset`` keeps that order continuous when
    a job comes back split across several JSONL lines.
    """
    result = OcrResult()
    pages = (body.get("result") or {}).get("layoutParsingResults") or []

    for page_no, page in enumerate(pages, start=page_offset):
        md = page.get("markdown") or {}
        if isinstance(md, dict) and md.get("text"):
            result.markdown += md["text"] + "\n"

        images = {k: v for k, v in (md.get("images") or {}).items()} if isinstance(md, dict) else {}
        pruned = page.get("prunedResult") or {}

        last_rank = 0
        for item in pruned.get("parsing_res_list") or []:
            label = str(item.get("block_label") or "text")
            if label in SKIP:
                continue
            raw = item.get("block_order")
            # Tables come back with block_order = None. Treating that as 0 — as this
            # did — sorted every table to the *top* of its page, ahead of the title.
            # The list itself is in reading order, so an order-less block stays where
            # it was found: right after the one before it. Ranks are doubled so that
            # "right after" is still a whole number.
            rank = 2 * int(raw) if raw is not None else last_rank + 1
            last_rank = rank
            order = page_no * 10_000 + rank

            content = item.get("block_content")
            blob = None
            if label in {"image", "figure", "chart"} and isinstance(content, str) and content in images:
                blob = _image_bytes(images[content])

            result.blocks.append(
                OcrBlock(
                    label=label,
                    text="" if blob else str(content or ""),
                    order=order,
                    image=blob,
                    bbox=tuple(item["block_bbox"]) if item.get("block_bbox") else None,
                )
            )

    result.blocks.sort(key=lambda b: b.order)
    return result


def parse_jsonl(lines: list[str]) -> OcrResult:
    """Merge the result JSONL into one reading-ordered document.

    A job's pages can arrive spread over several lines, so the page counter has to run
    across them: restarting it per line would interleave page 2 back into page 1.
    """
    merged = OcrResult()
    page_offset = 0
    for line in lines:
        body = json.loads(line)
        part = parse_response(body, page_offset)
        merged.blocks.extend(part.blocks)
        merged.markdown += part.markdown
        page_offset += len((body.get("result") or {}).get("layoutParsingResults") or [])
    merged.blocks.sort(key=lambda b: b.order)
    return merged


TEMPLATE = template_dir() / "student_reference.docx"


def blocks_to_docx(result: OcrResult, out: Path) -> Path:
    """Typeset OCR blocks into a Word file the rest of the pipeline can consume.

    The body is emptied through ``blank_template`` rather than by deleting
    ``doc.paragraphs``: the template ships with a demo *table* as well as demo
    paragraphs, and dropping only the paragraphs left that "Table 1 2" grid sitting at
    the top of every OCR'd paper — the same bug blank_template was written to fix for
    the answer sheet. It also pins A4 rather than inheriting US Letter.
    """
    doc = ds.blank_template(TEMPLATE)
    ds.ensure_note_styles(doc)

    import io

    for block in result.blocks:
        if block.image:
            try:
                doc.add_picture(io.BytesIO(block.image), width=Inches(4.5))
                continue
            except Exception:
                continue  # an undecodable figure must not sink the whole paper

        text = ds.sanitize(block.text).strip()
        if not text:
            continue
        if block.label == "table" and "<" in text:
            _add_table(doc, text)
            continue
        style = ds.NOTE_HEADING if block.label in HEADINGS else ds.NOTE_BODY
        for line in text.split("\n"):
            if line.strip():
                doc.add_paragraph(line, style=style)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def _add_table(doc, html: str) -> None:
    """Render the HTML table PaddleOCR returns for table blocks."""
    from html.parser import HTMLParser

    class Rows(HTMLParser):
        def __init__(self):
            super().__init__()
            self.rows: list[list[str]] = []
            self.cell: str | None = None

        def handle_starttag(self, tag, attrs):
            if tag == "tr":
                self.rows.append([])
            elif tag in ("td", "th"):
                self.cell = ""

        def handle_data(self, data):
            if self.cell is not None:
                self.cell += data

        def handle_endtag(self, tag):
            if tag in ("td", "th") and self.cell is not None and self.rows:
                self.rows[-1].append(self.cell.strip())
                self.cell = None

    parser = Rows()
    parser.feed(html)
    rows = [r for r in parser.rows if r]
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    ds.set_table_borders(table)
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.cell(i, j).text = ds.sanitize(cell)


def ingest_pdf(pdf: Path, out_dir: Path, *, base_url: str = "", token: str = "",
               sleep=time.sleep, log=None) -> Path:
    """OCR a PDF and write a .docx next to the other input papers.

    The service address is no longer something the teacher has to find: v2 is one
    global endpoint, so only the token is required.
    """
    token = token or os.environ.get(TOKEN_ENV, "")
    base_url = base_url or os.environ.get(BASE_URL_ENV, "")
    if not token:
        raise SystemExit(
            f"缺少 PaddleOCR access token。请在 https://aistudio.baidu.com/account/accessToken "
            f"获取后设置环境变量 {TOKEN_ENV}。"
        )

    jobs_url = _jobs_url(base_url)
    job_id = _submit(pdf, token, jobs_url)
    jsonl_url = _wait(job_id, token, jobs_url, sleep=sleep, log=log)
    result = parse_jsonl(_download_jsonl(jsonl_url))
    if not result.blocks:
        raise SystemExit(f"{pdf.name}: PaddleOCR 未识别出任何内容。")
    return blocks_to_docx(result, out_dir / f"{pdf.stem}.docx")
