#!/usr/bin/env python3
"""Generate the three committed Pandoc reference DOCX templates.

Requires python-docx only when regenerating templates. Runtime export remains
stdlib + Pandoc and consumes the generated assets directly.
"""

from __future__ import annotations

import io
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "assets" / "word_templates"
BLUE = RGBColor(24, 73, 132)
DARK = RGBColor(31, 41, 55)
MUTED = RGBColor(75, 85, 99)


def default_reference_bytes() -> bytes:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        raise RuntimeError("Pandoc is required to build reference templates")
    result = subprocess.run(
        [pandoc, "--print-default-data-file", "reference.docx"],
        capture_output=True,
        check=True,
    )
    return result.stdout


def set_fonts(style, *, latin: str, east_asia: str, size: float, bold: bool | None = None) -> None:
    font = style.font
    font.name = latin
    font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_run_fonts(run, *, latin: str, east_asia: str) -> None:
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_keep_with_next(style) -> None:
    ppr = style.element.get_or_add_pPr()
    if ppr.find(qn("w:keepNext")) is None:
        ppr.append(OxmlElement("w:keepNext"))


def set_page_break_before(style, enabled: bool) -> None:
    ppr = style.element.get_or_add_pPr()
    existing = ppr.find(qn("w:pageBreakBefore"))
    if existing is not None:
        ppr.remove(existing)
    if enabled:
        ppr.append(OxmlElement("w:pageBreakBefore"))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def configure_template(kind: str) -> Document:
    doc = Document(io.BytesIO(default_reference_bytes()))
    body_size = {"student": 11.0, "teacher": 10.5, "answers": 10.0}[kind]
    margins = {"student": 2.0, "teacher": 2.2, "answers": 1.8}[kind]
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(margins)
        section.bottom_margin = Cm(margins)
        section.left_margin = Cm(margins)
        section.right_margin = Cm(margins)
        section.header_distance = Cm(0.9)
        section.footer_distance = Cm(0.9)
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run({
            "student": "高三英语精选训练",
            "teacher": "高三英语教师讲解",
            "answers": "高三英语答案汇总",
        }[kind])
        set_run_fonts(hr, latin="Arial", east_asia="Arial Unicode MS")
        hr.font.size = Pt(8.5)
        hr.font.color.rgb = MUTED
        section.footer.is_linked_to_previous = False
        fp = section.footer.paragraphs[0]
        add_page_number(fp)

    normal = doc.styles["Normal"]
    set_fonts(normal, latin="Times New Roman", east_asia="Arial Unicode MS", size=body_size)
    normal.paragraph_format.line_spacing = {"student": 1.35, "teacher": 1.28, "answers": 1.15}[kind]
    normal.paragraph_format.space_after = Pt({"student": 5, "teacher": 4, "answers": 2}[kind])
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Title": (20, BLUE),
        "Heading 1": (18, BLUE),
        "Heading 2": (15, BLUE),
        "Heading 3": (12.5, DARK),
        "Heading 4": (11, BLUE),
    }
    for name, (size, color) in heading_specs.items():
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        set_fonts(style, latin="Arial", east_asia="Arial Unicode MS", size=size, bold=True)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10 if name != "Heading 1" else 0)
        style.paragraph_format.space_after = Pt(6)
        set_keep_with_next(style)
    if "Heading 1" in doc.styles:
        doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.styles["Heading 1"].paragraph_format.space_before = Pt(150 if kind in {"student", "teacher"} else 0)
        doc.styles["Heading 1"].paragraph_format.space_after = Pt(20)
    if "Heading 2" in doc.styles:
        set_page_break_before(doc.styles["Heading 2"], kind in {"student", "teacher"})

    for name in ("List Bullet", "List Number"):
        if name in doc.styles:
            set_fonts(doc.styles[name], latin="Times New Roman", east_asia="Arial Unicode MS", size=body_size)
            doc.styles[name].paragraph_format.left_indent = Cm(0.65)
            doc.styles[name].paragraph_format.first_line_indent = Cm(-0.3)
            doc.styles[name].paragraph_format.space_after = Pt(2.5)
    if "Block Text" in doc.styles:
        set_fonts(doc.styles["Block Text"], latin="Times New Roman", east_asia="Arial Unicode MS", size=body_size)
    if "Source Code" in doc.styles:
        set_fonts(doc.styles["Source Code"], latin="Consolas", east_asia="Arial Unicode MS", size=max(9, body_size - 1))
    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for kind, filename in [
        ("student", "student_reference.docx"),
        ("teacher", "teacher_reference.docx"),
        ("answers", "answers_reference.docx"),
    ]:
        path = OUT_DIR / filename
        configure_template(kind).save(path)
        print(f"wrote {path}")


if __name__ == "__main__":
    main()
