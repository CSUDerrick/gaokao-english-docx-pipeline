#!/usr/bin/env python3
"""Build the student and teacher Word files by cloning the source papers.

Replaces the Markdown -> Pandoc path for the two documents that contain question
text. Going through Markdown meant the original typesetting had to be re-invented
(and images were dropped entirely); here each question's original ``w:p``/``w:tbl``
nodes are copied straight out of the paper it came from.

The answers-only document has no source typesetting to preserve — it is answer
letters and model essays — so it is built from the structured data instead.
"""

from __future__ import annotations

import json
import shutil
import tempfile
from pathlib import Path

import docx_splice as ds
from bundle_paths import template_dir

STUDENT = "高三英语精选试题_学生版.docx"
TEACHER = "高三英语精选试题_教师讲解版.docx"
ANSWERS = "高三英语精选试题_答案汇总版.docx"


def _load(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def _notes_for(row: dict, pipeline) -> list[tuple[str, str]]:
    score = row.get("score") if isinstance(row.get("score"), dict) else {}
    enrichment = row.get("enrichment") if isinstance(row.get("enrichment"), dict) else {}
    s = (score or {}) | (enrichment or {})

    notes: list[tuple[str, str]] = [
        (ds.NOTE_HEADING, "教师讲解"),
        (
            ds.NOTE_BODY,
            f"主题：{s.get('topic', '')}\n"
            f"评分：新颖度 {s.get('novelty_score', '')}；难度 {s.get('difficulty_score', '')}；"
            f"词汇价值 {s.get('vocabulary_value_score', '')}；语法价值 {s.get('grammar_value_score', '')}；"
            f"推荐度 {s.get('recommendation_score', '')}\n"
            f"入选理由：{s.get('selection_reason', '')}\n"
            f"课堂建议：{s.get('classroom_suggestion', '')}",
        ),
    ]

    sections = [
        ("重点词汇", "\n".join(
            x for x in (
                pipeline.render_words(s.get("core_high_frequency_words", [])),
                pipeline.render_words(s.get("familiar_words_new_meanings", [])),
                pipeline.render_words(s.get("difficult_or_low_frequency_words", [])),
            ) if x.strip()
        )),
        ("语法与词汇变形", pipeline.render_grammar(s.get("word_formation_and_grammar", []))),
        ("长难句", pipeline.render_sentences(s.get("long_difficult_sentences", []))),
        ("补充讲解建议", str(s.get("teaching_notes", "") or "暂无")),
    ]
    for title, body in sections:
        if body and body.strip():
            notes.append((ds.NOTE_HEADING, title))
            notes.append((ds.NOTE_BODY, body))
    return notes


TEMPLATE_DIR = template_dir()


def _force_a4(doc) -> None:
    """Set A4 explicitly rather than trusting the template to be found.

    python-docx defaults to US Letter, so if the template ever goes missing the
    answer sheet silently comes out the wrong size — which is exactly what
    happened inside the packaged app.
    """
    from docx.shared import Cm

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)


def _answers_doc(rows: list[dict], out: Path, pipeline) -> Path:
    from docx import Document

    # Built on the existing A4 template so the answer sheet keeps the Chinese
    # font mapping and footer that the student/teacher versions inherit from the
    # source papers.
    template = TEMPLATE_DIR / "answers_reference.docx"
    doc = Document(str(template)) if template.exists() else Document()
    for para in list(doc.paragraphs):
        para._element.getparent().remove(para._element)
    _force_a4(doc)

    ds.ensure_note_styles(doc)
    doc.add_paragraph("高三英语答案汇总", style=ds.NOTE_HEADING)

    current = ""
    for row in rows:
        section = row.get("display_section") or pipeline.section_display(row.get("section", ""))
        if section != current:
            doc.add_paragraph(section, style=ds.NOTE_HEADING)
            current = section
        segment = _load(Path(row["segment_path"]))
        doc.add_paragraph(str(row.get("item_label", "")), style=ds.NOTE_ANSWER)
        doc.add_paragraph(pipeline.answer_key_text(segment.get("answer_key")), style=ds.NOTE_BODY)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def _plain_doc(heading: str, text: str, notes: list[tuple[str, str]], out: Path) -> Path:
    """Fallback for a question that could not be traced back to its source.

    Loses the original typesetting, but keeps the question in the paper.
    """
    from docx import Document

    template = TEMPLATE_DIR / "student_reference.docx"
    doc = Document(str(template)) if template.exists() else Document()
    for para in list(doc.paragraphs):
        para._element.getparent().remove(para._element)
    _force_a4(doc)

    ds.ensure_note_styles(doc)
    doc.add_paragraph(ds.sanitize(heading), style=ds.NOTE_HEADING)
    for line in ds.sanitize(text).split("\n"):
        if line.strip():
            doc.add_paragraph(line, style=ds.NOTE_BODY)
    for style_name, body in notes:
        for line in ds.sanitize(body).split("\n"):
            if line.strip():
                doc.add_paragraph(line, style=style_name)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def export_selected(out_dir: Path, pipeline, log=print) -> list[Path]:
    selection = out_dir / "selected_items.json"
    if not selection.exists():
        raise SystemExit(f"Missing {selection}. Run --mode select first.")
    rows = _load(selection)

    enriched = out_dir / "selected_items.enriched.json"
    if enriched.exists():
        pipeline.merge_cached_enrichments(rows, _load(enriched))

    rows = sorted(
        rows,
        key=lambda r: (pipeline.section_order(r.get("section", "")), r.get("source_doc", ""), r.get("item_id", "")),
    )

    docx_dir = out_dir / "docx_exports"
    docx_dir.mkdir(parents=True, exist_ok=True)

    student_parts: list[Path] = []
    teacher_parts: list[Path] = []
    skipped: list[str] = []

    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        for n, row in enumerate(rows):
            segment = _load(Path(row["segment_path"]))
            src = segment.get("source_path")
            blocks = segment.get("source_blocks") or []
            heading = f"{row.get('item_label', '')}｜来源：{row.get('source_doc', '')}"
            cloneable = bool(src) and len(blocks) == 2 and Path(src).exists()

            if not cloneable:
                # No pointer back into an original docx (AI-segmented text the
                # anchor could not place). Render the text we do have rather than
                # dropping the question — a missing question is far worse than an
                # unstyled one.
                skipped.append(str(row.get("item_id")))
                text = segment.get("question_text") or ""
                sp = _plain_doc(heading, text, [], tmp / f"s{n:03d}.docx")
                tp = _plain_doc(heading, text, _notes_for(row, pipeline), tmp / f"t{n:03d}.docx")
                student_parts.append(sp)
                teacher_parts.append(tp)
                continue

            lo, hi = blocks
            indices = list(range(lo, hi))

            sp = tmp / f"s{n:03d}.docx"
            ds.clone_subset(Path(src), indices, sp)
            ds.decorate(sp, heading=heading)
            student_parts.append(sp)

            tp = tmp / f"t{n:03d}.docx"
            ds.clone_subset(Path(src), indices, tp)
            ds.decorate(tp, heading=heading, notes=_notes_for(row, pipeline))
            teacher_parts.append(tp)

        if skipped:
            log(f"  WARNING: {len(skipped)} 道题无法定位到原卷段落，已按纯文本导出（格式需手动调整）：{', '.join(skipped)}")
        if not student_parts:
            raise SystemExit("No selected item produced any content; nothing to export.")

        created: list[Path] = []
        for parts, name, title in (
            (student_parts, STUDENT, "高三英语精选试题（学生版）"),
            (teacher_parts, TEACHER, "高三英语精选试题（教师讲解版）"),
        ):
            target = docx_dir / name
            ds.merge(parts, target)
            ds.scrub_metadata(target, title)
            created.append(target)
            log(f"  wrote {target.name}  (images: {ds.media_count(target)})")

    answers = _answers_doc(rows, docx_dir / ANSWERS, pipeline)
    ds.scrub_metadata(answers, "高三英语答案汇总")
    log(f"  wrote {answers.name}")
    created.append(answers)

    for path in created:
        ds.validate(path)

    return created
