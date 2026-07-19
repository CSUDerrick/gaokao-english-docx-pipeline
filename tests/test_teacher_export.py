"""The teacher edition is 原题 + 官方答案与解析 + 详细解析和解答步骤, in that order.

The regression that matters most here is the *student* edition: this change rebuilt
the teacher path, and the student paper is the one that gets handed to a room full
of teenagers. So these tests also pin that the student file still carries no source
attribution, no answer key and no explanation.
"""

from __future__ import annotations

import json
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import docx_blocks as db  # noqa: E402
import export_docx_splice as ex  # noqa: E402
import gaokao_english_docx_pipeline as pipeline  # noqa: E402
from docx import Document  # noqa: E402
from docx.shared import Mm  # noqa: E402

QUESTION = [
    "B",
    "Cui Yameng runs a barrier-free guesthouse in her home city.",
    "24. What inspired Cui to design the guesthouse?",
    "A. A news report.  B. A hotel stay.  C. Her travels.  D. A friend's advice.",
    "25. What does the underlined word mean?",
    "A. Refused.  B. Encouraged.  C. Delayed.  D. Ignored.",
]
ANSWERS = [
    "参考答案",
    "24—25. CB",
    "【导语】本文主要讲述崔雅梦创办无障碍民宿帮助残障人士的故事。",
    "【24题详解】",
    "细节理解题。根据第二段可知，灵感来源于她与残障朋友的旅行。",
    "【25题详解】",
    "词义猜测题。根据上下文可知，该词意为“鼓舞”。",
]

EXPLANATION = {
    "questions": [
        {
            "number": "24",
            "answer": "C",
            "question_type": "细节理解题",
            "locate": "第二段：\"Her travels with physically-challenged friends…\"",
            "reasoning": "原文的 travels 被选项换成了 Her travels，属于同义替换。",
            "distractors": [{"option": "B", "why_wrong": "文中确实提到酒店，但那是她的目标，不是灵感来源。"}],
            "language_note": "",
        },
        {
            "number": "25",
            "answer": "B",
            "question_type": "词义猜测题",
            "locate": "第三段：\"The post inspired many followers…\"",
            "reasoning": "后一句举了一个受鼓舞出行的例子，方向是正面的。",
            "distractors": [{"option": "A", "why_wrong": "Refused 是反方向，和后文的 embarked on his first trip 冲突。"}],
        },
    ]
}


def _paper(tmp: Path) -> Path:
    doc = Document()
    # Real papers are A4, and validate() refuses anything else — python-docx would
    # otherwise hand us US Letter and fail the export on the fixture, not the code.
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    for line in QUESTION:
        doc.add_paragraph(line)
    for line in ANSWERS:
        doc.add_paragraph(line)
    path = tmp / "湖北某中学2026届高三英语试题.docx"
    doc.save(str(path))
    return path


def _run_dir(tmp: Path) -> Path:
    """A finished run: one selected item, with its segment and its explanation."""
    src = _paper(tmp)
    doc = db.read_docx(src)
    question_end = doc.text.index("参考答案")
    lo, hi = doc.body_range(0, question_end)

    import answer_explanation as ax

    official = ax.OfficialExplanations(doc, question_end)
    explanation_blocks = official.blocks_for([24, 25])
    assert explanation_blocks, "fixture must have an official explanation to clone"

    out = tmp / "run"
    (out / "segments").mkdir(parents=True)
    segment = {
        "item_id": "湖北__reading_b__01",
        "source_doc": src.name,
        "section": "reading_b",
        "display_section": "阅读B",
        "item_label": "阅读B",
        "question_text": "\n".join(QUESTION),
        "answer_key": [{"number": "24", "answer": "C"}, {"number": "25", "answer": "B"}],
        "answer_source": "答案区",
        "source_path": str(src),
        "source_blocks": [lo, hi],
        "official_explanation_blocks": explanation_blocks,
    }
    segment_path = out / "segments" / "item.json"
    segment_path.write_text(json.dumps(segment, ensure_ascii=False), encoding="utf-8")

    rows = [{
        "item_id": "湖北__reading_b__01",
        "source_doc": src.name,
        "section": "reading_b",
        "display_section": "阅读B",
        "item_label": "阅读B",
        "segment_path": str(segment_path),
        "explanation": EXPLANATION,
        "has_official_explanation": True,
    }]
    (out / "selected_items.json").write_text(json.dumps(rows, ensure_ascii=False), encoding="utf-8")
    return out


def _text(path: Path) -> str:
    return db.read_docx(path).text


def test_teacher_edition_has_question_then_official_then_ai_in_that_order():
    with tempfile.TemporaryDirectory() as tmp:
        out = _run_dir(Path(tmp))
        ex.export_selected(out, pipeline, log=lambda *_: None)
        text = _text(out / "docx_exports" / ex.TEACHER)

        question = text.index("What inspired Cui to design")
        official = text.index(ex.OFFICIAL_HEADING)
        detail = text.index("【24题详解】")
        ours = text.index(ex.AI_HEADING)
        assert question < official < detail < ours, "原题 → 官方答案与解析 → 详细解析和解答步骤"


def test_teacher_edition_carries_the_papers_own_explanation_verbatim():
    with tempfile.TemporaryDirectory() as tmp:
        out = _run_dir(Path(tmp))
        ex.export_selected(out, pipeline, log=lambda *_: None)
        text = _text(out / "docx_exports" / ex.TEACHER)

        assert "【导语】本文主要讲述崔雅梦创办无障碍民宿" in text
        assert "细节理解题。根据第二段可知" in text
        assert "【25题详解】" in text


def test_teacher_edition_carries_our_explanation_including_the_wrong_options():
    with tempfile.TemporaryDirectory() as tmp:
        out = _run_dir(Path(tmp))
        ex.export_selected(out, pipeline, log=lambda *_: None)
        text = _text(out / "docx_exports" / ex.TEACHER)

        assert "24. C　细节理解题" in text
        assert "为什么不选 B" in text, "the distractor analysis is the point of the AI pass"
        assert "同义替换" in text


def test_a_paper_that_never_explained_the_question_says_so_and_still_prints_the_answers():
    with tempfile.TemporaryDirectory() as tmp:
        out = _run_dir(Path(tmp))
        segment_path = out / "segments" / "item.json"
        segment = json.loads(segment_path.read_text(encoding="utf-8"))
        segment["official_explanation_blocks"] = []  # 广东 skips four whole sections
        segment_path.write_text(json.dumps(segment, ensure_ascii=False), encoding="utf-8")

        ex.export_selected(out, pipeline, log=lambda *_: None)
        text = _text(out / "docx_exports" / ex.TEACHER)

        assert ex.NO_OFFICIAL in text
        assert "24: C" in text, "the answer key is known even when the explanation is not"
        assert ex.AI_HEADING in text, "our explanation is written either way"


def _body_sz_values(path: Path) -> set[str]:
    import zipfile

    from lxml import etree

    w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    with zipfile.ZipFile(path) as z:
        dx = etree.fromstring(z.read("word/document.xml"))
    return {e.get(w + "val") for e in dx.findall(".//" + w + "sz")}


def test_every_exported_file_shares_one_house_style():
    """统一格式：the student, teacher and answer files all come out 小四, TNR + 宋体,
    single-spaced, snap-to-grid off — regardless of what font the source paper used."""
    import zipfile

    from lxml import etree

    w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    with tempfile.TemporaryDirectory() as tmp:
        out = _run_dir(Path(tmp))
        ex.export_selected(out, pipeline, log=lambda *_: None)

        for name in (ex.STUDENT, ex.TEACHER, ex.ANSWERS):
            path = out / "docx_exports" / name
            assert _body_sz_values(path) == {"24"}, f"{name} 不是统一小四: {_body_sz_values(path)}"
            with zipfile.ZipFile(path) as z:
                dx = etree.fromstring(z.read("word/document.xml"))
            assert dx.findall(".//" + w + "contextualSpacing") == [], f"{name} 仍勾了「相同样式不加空格」"
            for p in dx.findall(".//" + w + "p"):
                snap = p.find(w + "pPr/" + w + "snapToGrid")
                assert snap is not None and snap.get(w + "val") == "0", f"{name} 有段落仍对齐网格"


def test_missing_question_numbers_flags_a_dropped_reading_passage():
    import segment_quality as sq

    complete = {
        "section": "reading_b",
        "question_text": "B\nA passage.\n24. Q?\nA. x B. y\n25. Q?\nA. x B. y",
        "answer_key": [{"number": "24"}, {"number": "25"}],
    }
    dropped = {
        "section": "reading_b",
        "question_text": "B\nA passage with no questions at all.",
        "answer_key": [{"number": "24"}, {"number": "25"}, {"number": "26"}, {"number": "27"}],
    }
    cloze = {  # 完形 numbers blanks inline, not as "N." stems — must not cry wolf
        "section": "cloze",
        "question_text": "A passage with ____41____ and ____42____ blanks.",
        "answer_key": [{"number": "41"}, {"number": "42"}],
    }
    assert sq.missing_question_numbers(complete) == []
    assert sq.missing_question_numbers(dropped) == ["24", "25", "26", "27"]
    assert sq.missing_question_numbers(cloze) == []


def test_a_reading_passage_missing_its_questions_is_skipped_not_shipped():
    """One 华南师范 卷 had passage B's questions overwritten by passage C's (决策 35).

    The stems are simply not in the source, so the passage cannot be exported into
    anything a student can answer. It must be dropped with a warning rather than
    shipped as a passage with no questions — and the complete questions stay."""
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        out = _run_dir(tmp)

        broken = {
            "item_id": "缺题卷__reading_c__01",
            "source_doc": "缺题卷2026届高三英语试题.docx",
            "section": "reading_c",
            "display_section": "阅读C",
            "item_label": "阅读C",
            "question_text": "C\nDEFECTIVE_PASSAGE_MARKER a passage whose questions the paper dropped.",
            "answer_key": [{"number": "28"}, {"number": "29"}, {"number": "30"}, {"number": "31"}],
            "source_path": "",
            "source_blocks": None,
        }
        broken_path = out / "segments" / "broken.json"
        broken_path.write_text(json.dumps(broken, ensure_ascii=False), encoding="utf-8")

        rows = json.loads((out / "selected_items.json").read_text(encoding="utf-8"))
        rows.append({
            "item_id": "缺题卷__reading_c__01",
            "source_doc": "缺题卷2026届高三英语试题.docx",
            "section": "reading_c",
            "display_section": "阅读C",
            "item_label": "阅读C",
            "segment_path": str(broken_path),
        })
        (out / "selected_items.json").write_text(json.dumps(rows, ensure_ascii=False), encoding="utf-8")

        logs: list[str] = []
        ex.export_selected(out, pipeline, log=logs.append)

        for name in (ex.STUDENT, ex.TEACHER):
            text = _text(out / "docx_exports" / name)
            assert "DEFECTIVE_PASSAGE_MARKER" not in text, f"{name} 混入了缺题的阅读段"
            assert "What inspired Cui to design" in text, f"{name} 丢了完整的题"
        assert any("缺第" in line and "28" in line for line in logs), "缺题应有警告"


def test_the_student_edition_is_untouched_by_any_of_this():
    with tempfile.TemporaryDirectory() as tmp:
        out = _run_dir(Path(tmp))
        ex.export_selected(out, pipeline, log=lambda *_: None)
        text = _text(out / "docx_exports" / ex.STUDENT)

        assert "What inspired Cui to design" in text, "the question is still there"
        assert "第 1 篇" in text
        for forbidden in ("【24题详解】", "【导语】", ex.OFFICIAL_HEADING, ex.AI_HEADING, "来源：", "24—25. CB"):
            assert forbidden not in text, f"学生版混入了 {forbidden}"
        # And the gate that would have stopped the file shipping at all.
        ex.assert_student_edition_is_clean(out / "docx_exports" / ex.STUDENT)
