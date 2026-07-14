"""Boundary repair: the model reports *where*, never *what*.

A real paper headed its 七选五 with only ``第二节（共5小题...）`` — no "七选五", no
"多余选项" — so the keyword rules missed it and 阅读D swallowed it (4,722 chars vs
~2,800 for its neighbours). That particular paper is now handled locally, by
matching its mark scheme (5 questions, 12.5 marks, unique in a paper) — see
``test_mark_scheme_heading_needs_no_repair``.

Repair still exists for papers where no rule can work at all, and these tests use
a 七选五 headed by a bare 第二节 to exercise it.

The old fallback handed the whole paper to the model and asked it to re-emit every
question. It came back with 44 items, several 37 characters long, none of which
could be traced back to the original paragraphs — so the export could not clone
the teacher's formatting.

Now the model is asked for a single paragraph number and the text is still cut
from the original document.
"""

from __future__ import annotations

import importlib.util
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import docx_blocks as db  # noqa: E402
import segment_repair as sr  # noqa: E402
from segment_quality import evaluate_document  # noqa: E402

_spec = importlib.util.spec_from_file_location("pipeline_repair", ROOT / "scripts" / "gaokao_english_docx_pipeline.py")
pipeline = importlib.util.module_from_spec(_spec)
sys.modules["pipeline_repair"] = pipeline
_spec.loader.exec_module(pipeline)


def _paper(path: Path, gap_heading: str = "第二节") -> Path:
    """A paper whose 七选五 carries no recognisable name.

    ``gap_heading`` defaults to a bare 第二节: no "七选五", no "多余选项", and no
    mark scheme either, so no local rule can find it and the model has to be
    asked where the boundary is. Pass the 江苏 heading to get the paper that the
    local rules *do* handle now (see test_mark_scheme_heading_needs_no_repair).
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("第二部分 阅读理解")
    doc.add_paragraph("A")
    doc.add_paragraph("Passage one talks about a school and its teaching methods in detail.")
    for n in (21, 22, 23):
        doc.add_paragraph(f"{n}. Question about passage A here?")
        doc.add_paragraph("A. one B. two C. three D. four")

    doc.add_paragraph("D")
    doc.add_paragraph("Grief can bring waves of heartache but for most people those feelings ease.")
    for n in (32, 33, 34, 35):
        doc.add_paragraph(f"{n}. Question about passage D here?")
        doc.add_paragraph("A. one B. two C. three D. four")

    # 七选五 — headed only by `gap_heading`
    doc.add_paragraph(gap_heading)
    doc.add_paragraph("As a child growing up there are many different paths available. ____36____")
    doc.add_paragraph("I was always looking for pennies. ____37____ My father passed away. ____38____")
    doc.add_paragraph("Later I told my neighbor the story. ____39____ How things change. ____40____")
    doc.add_paragraph("A. He passed the habit on to me.")
    doc.add_paragraph("B. I figured he would not think more of it.")
    doc.add_paragraph("G. His soft words supported me through my youth.")

    doc.add_paragraph("第三部分 语言知识运用")
    doc.add_paragraph("完形填空")
    doc.add_paragraph("阅读下面短文，从各题所给的四个选项中选出最佳选项。" + "填空正文。" * 40)

    doc.add_paragraph("答案")
    doc.add_paragraph("21-23. BCD")
    doc.add_paragraph("36-40. ABGCF")
    doc.save(str(path))
    return path


def test_mark_scheme_heading_needs_no_repair():
    """The 江苏 paper: 七选五 named nowhere, but its mark scheme gives it away.

    This used to cost a model round-trip on every run. A section worth 12.5 marks
    over 5 questions is 七选五 and nothing else, so the rules can place it locally.
    """
    with tempfile.TemporaryDirectory() as tmp:
        src = _paper(Path(tmp) / "a.docx", gap_heading="第二节（共5小题；每小题2.5分，满分12.5分）")
        doc = db.read_docx(src)
        segments = pipeline.local_segment_paper(src.name, doc.text, doc)

        gap = [s for s in segments if s["section"] == "gap_filling"]
        assert gap, "七选五 must be found without asking the model"
        assert "____36____" in gap[0]["question_text"]
        # ...and it must no longer be hiding inside 阅读D.
        reading_d = [s for s in segments if s["section"] == "reading_d"][0]
        assert "____36____" not in reading_d["question_text"]


def test_host_range_points_at_the_section_that_swallowed_it():
    # 七选五 must be hiding inside 阅读D, the section before it in canonical order.
    segments = [
        {"section": "reading_a", "source_blocks": [10, 20]},
        {"section": "reading_d", "source_blocks": [20, 60]},
        {"section": "cloze", "source_blocks": [60, 80]},
    ]
    assert sr._host_range(segments, "gap_filling") == (20, 60)


def test_prompt_covers_only_the_suspect_region_not_the_whole_paper():
    with tempfile.TemporaryDirectory() as tmp:
        doc = db.read_docx(_paper(Path(tmp) / "a.docx"))
        whole = len(doc.text)
        prompt = sr.build_prompt(doc, 8, 20, "gap_filling")
        assert len(prompt) < whole, "the point is to send a slice, not the entire paper"
        assert "[8]" in prompt or "[9]" in prompt
        assert "七选五" in prompt


def test_reply_parsing_accepts_the_shapes_a_model_actually_returns():
    assert sr.parse_start_block('{"start_block": 135}', 114, 149) == 135
    assert sr.parse_start_block('```json\n{"start_block":135}\n```', 114, 149) == 135
    assert sr.parse_start_block("135", 114, 149) == 135


def test_out_of_range_or_absent_answers_fail_closed():
    # Applying a hallucinated index would cut the paper in the wrong place.
    assert sr.parse_start_block('{"start_block": 9999}', 114, 149) is None
    assert sr.parse_start_block('{"start_block": null}', 114, 149) is None
    assert sr.parse_start_block("", 114, 149) is None
    assert sr.parse_start_block('{"start_block": 114}', 114, 149) is None, "equal to lo is not a split"
    assert sr.parse_start_block('{"start_block": 200}', 114, 149) is None


def test_repair_recovers_the_missing_section_and_keeps_it_cloneable():
    with tempfile.TemporaryDirectory() as tmp:
        src = _paper(Path(tmp) / "a.docx")
        doc = db.read_docx(src)

        before = pipeline.local_segment_paper(src.name, doc.text, doc)
        sections = {s["section"] for s in before}
        assert "gap_filling" not in sections, "fixture must reproduce the miss"

        start = [b.body_index for b in doc.blocks if b.text.strip() == "第二节"][0]
        extra = sr.locate_missing_sections(
            doc, before, ["gap_filling"],
            ask=lambda _p: f'{{"start_block": {start}}}',
            log=lambda _m: None,
        )
        assert extra, "the boundary should have been located"

        after = pipeline.local_segment_paper(src.name, doc.text, doc, extra_starts=extra)
        gap = [s for s in after if s["section"] == "gap_filling"]
        assert gap, "七选五 must now exist"

        seg = gap[0]
        assert "____36____" in seg["question_text"]
        assert seg["source_blocks"], "must stay cloneable — otherwise formatting is lost"

        # the text must come from the ORIGINAL document, not from the model
        lo, hi = seg["source_blocks"]
        original = "\n".join(b.text for b in doc.blocks if lo <= b.body_index < hi)
        assert seg["question_text"].strip() in original

        # and the section it was hiding in must have shrunk accordingly
        d_before = [s for s in before if s["section"] == "reading_d"][0]
        d_after = [s for s in after if s["section"] == "reading_d"][0]
        assert len(d_after["question_text"]) < len(d_before["question_text"])


def test_repair_turns_the_warning_into_a_pass():
    with tempfile.TemporaryDirectory() as tmp:
        src = _paper(Path(tmp) / "a.docx")
        doc = db.read_docx(src)
        before = pipeline.local_segment_paper(src.name, doc.text, doc)

        def rows_for(segs):
            return [
                {"source_doc": src.name, "section": s["section"], "item_label": s["item_label"],
                 "char_count": len(s["question_text"]),
                 "answer_count": len(s["answer_key"]) if isinstance(s["answer_key"], list) else 0}
                for s in segs
            ]

        assert "gap_filling" in evaluate_document(src.name, rows_for(before))["missing"]

        start = [b.body_index for b in doc.blocks if b.text.strip() == "第二节"][0]
        extra = sr.locate_missing_sections(
            doc, before, ["gap_filling"],
            ask=lambda _p: f'{{"start_block": {start}}}', log=lambda _m: None,
        )
        after = pipeline.local_segment_paper(src.name, doc.text, doc, extra_starts=extra)
        assert "gap_filling" not in evaluate_document(src.name, rows_for(after))["missing"]


def test_model_declining_leaves_the_local_result_alone():
    with tempfile.TemporaryDirectory() as tmp:
        src = _paper(Path(tmp) / "a.docx")
        doc = db.read_docx(src)
        before = pipeline.local_segment_paper(src.name, doc.text, doc)

        extra = sr.locate_missing_sections(
            doc, before, ["gap_filling"],
            ask=lambda _p: '{"start_block": null}', log=lambda _m: None,
        )
        assert extra == [], "no boundary means no change, not a bad guess"
