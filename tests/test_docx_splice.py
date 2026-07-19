"""Fidelity tests for the clone-and-splice export.

The bug these pin down: the old export read the source .docx into a flat string,
so every bold run, underline (which is how a blank is drawn on an English paper),
table and image was gone before segmentation even started — the exported student
paper contained zero images and one 10,000-character paragraph.

These tests build their own .docx fixtures rather than reading ``input_docx/``,
which is gitignored, so they run on a fresh clone.
"""

from __future__ import annotations

import io
import sys
import tempfile
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import docx_blocks as db  # noqa: E402
import docx_splice as ds  # noqa: E402
from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.shared import Inches  # noqa: E402
from lxml import etree  # noqa: E402

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _png() -> io.BytesIO:
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (32, 32), (200, 60, 60)).save(buf, format="PNG")
    buf.seek(0)
    return buf


def _fixture(path: Path, marker: str = "X") -> Path:
    """A paper-ish docx: styled runs, a blank-line underline, a table, an image."""
    doc = Document()
    doc.add_paragraph(f"{marker} 阅读理解 Passage")

    p = doc.add_paragraph()
    p.add_run("The word ").bold = False
    p.add_run("crucial").bold = True
    p.add_run(" means ")
    p.add_run("        ").underline = True  # a fill-in blank
    p.add_run(" here.")

    p2 = doc.add_paragraph()
    p2.add_run("Emphasis").italic = True

    doc.add_paragraph("21. What does the author mean?  ____21____")

    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    t.cell(0, 0).text = "A"
    t.cell(0, 1).text = "B"
    t.cell(1, 0).text = "C"
    t.cell(1, 1).text = "D"

    doc.add_paragraph("Figure:")
    doc.add_picture(_png(), width=Inches(1))
    doc.add_paragraph("Tail paragraph")
    doc.save(str(path))
    return path


def _count(path: Path, tag: str) -> int:
    with zipfile.ZipFile(path) as z:
        dx = etree.fromstring(z.read("word/document.xml"))
    return len(dx.findall(".//" + W + tag))


def test_flat_text_drops_empty_paragraphs_and_joins_table_rows():
    with tempfile.TemporaryDirectory() as tmp:
        src = _fixture(Path(tmp) / "a.docx")
        doc = db.read_docx(src)
        assert "阅读理解 Passage" in doc.text
        assert "A | B" in doc.text and "C | D" in doc.text
        # every block's recorded span must actually index its own text
        for b in doc.blocks:
            assert doc.text[b.char_start : b.char_end] == b.text


def test_clone_preserves_bold_underline_table_and_image():
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        src = _fixture(tmp / "a.docx")
        doc = db.read_docx(src)

        out = ds.clone_subset(src, list(range(len(doc.body_children))), tmp / "clone.docx")

        for tag in ("b", "u", "i", "tbl", "drawing"):
            assert _count(out, tag) == _count(src, tag), f"<w:{tag}> not preserved"
        assert ds.media_count(out) == ds.media_count(src) == 1
        assert ds.unresolved_rids(out) == []


def test_clone_subset_keeps_only_requested_blocks():
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        src = _fixture(tmp / "a.docx")
        doc = db.read_docx(src)
        first = [b for b in doc.blocks if "Passage" in b.text][0]

        out = ds.clone_subset(src, [first.body_index], tmp / "one.docx")
        text = db.read_docx(out).text
        assert "Passage" in text
        assert "Tail paragraph" not in text


def test_clone_preserves_root_namespaces_and_mc_ignorable():
    # Re-serializing the root with ElementTree drops namespace declarations that
    # mc:Ignorable still names, which is what made Word report "unreadable
    # content". Every prefix listed in mc:Ignorable must stay declared.
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        src = _fixture(tmp / "a.docx")
        doc = db.read_docx(src)
        out = ds.clone_subset(src, list(range(len(doc.body_children))), tmp / "c.docx")

        with zipfile.ZipFile(out) as z:
            root = etree.fromstring(z.read("word/document.xml"))
        ignorable = (root.get("{http://schemas.openxmlformats.org/markup-compatibility/2006}Ignorable") or "").split()
        dangling = [p for p in ignorable if p not in root.nsmap]
        assert dangling == [], f"mc:Ignorable names undeclared prefixes: {dangling}"


def test_merge_across_documents_resolves_id_collisions():
    # Two independently authored Word files reuse the same styleIds and rIds for
    # different things. Merging must remap them, or a picture ends up pointing at
    # a footer and a table loses its borders.
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        a = _fixture(tmp / "a.docx", marker="A")
        b = _fixture(tmp / "b.docx", marker="B")
        pa = ds.clone_subset(a, list(range(len(db.read_docx(a).body_children))), tmp / "pa.docx")
        pb = ds.clone_subset(b, list(range(len(db.read_docx(b).body_children))), tmp / "pb.docx")

        out = ds.merge([pa, pb], tmp / "merged.docx")

        assert ds.unresolved_rids(out) == []
        assert _count(out, "drawing") == 2, "both images must survive the merge"
        assert ds.media_count(out) >= 1

        # every image relationship must point at a file that is actually present
        with zipfile.ZipFile(out) as z:
            doc_xml = z.read("word/document.xml").decode()
            rels = z.read("word/_rels/document.xml.rels").decode()
            media = {n for n in z.namelist() if n.startswith("word/media/")}
        import re

        targets = dict(re.findall(r'Id="([^"]+)"[^>]*Target="([^"]+)"', rels))
        for rid in set(re.findall(r'r:embed="([^"]+)"', doc_xml)):
            assert "word/" + targets[rid] in media, f"{rid} points at a missing part"

        text = db.read_docx(out).text
        assert "A 阅读理解" in text and "B 阅读理解" in text


def test_every_referenced_style_is_defined():
    # The historical corruption: referencing a style that styles.xml never
    # declares. Must hold after notes are added and documents merged.
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        src = _fixture(tmp / "a.docx")
        out = ds.clone_subset(src, list(range(len(db.read_docx(src).body_children))), tmp / "c.docx")
        ds.decorate(out, heading="阅读A｜来源：某卷", notes=[(ds.NOTE_BODY, "讲解正文")])

        with zipfile.ZipFile(out) as z:
            dx = etree.fromstring(z.read("word/document.xml"))
            st = etree.fromstring(z.read("word/styles.xml"))
        defined = {s.get(W + "styleId") for s in st.findall(W + "style")}
        used = {e.get(W + "val") for tag in ("pStyle", "rStyle", "tblStyle") for e in dx.findall(".//" + W + tag)}
        assert used <= defined, f"undefined styles referenced: {used - defined}"


def test_notes_strip_illegal_control_characters():
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        src = _fixture(tmp / "a.docx")
        out = ds.clone_subset(src, [0], tmp / "c.docx")
        ds.decorate(out, notes=[(ds.NOTE_BODY, "答案\x07：B\x00 blank ____36____")])

        with zipfile.ZipFile(out) as z:
            raw = z.read("word/document.xml")
        for ch in (b"\x00", b"\x07", b"\x0b", b"\x1f"):
            assert ch not in raw, "illegal XML control character reached the docx"
        # the blank marker must survive literally, not become bold (the old
        # Markdown path read ____ as __strong__)
        assert "____36____" in db.read_docx(out).text


def _bookmarked(path: Path) -> Path:
    """A paper whose bookmark spans two paragraphs — like Word's own OLE_LINKs."""
    from docx.shared import Cm

    doc = Document()
    for section in doc.sections:  # real papers are A4; validate() insists on it
        section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    doc.add_paragraph("第一段：题干开始。" * 6)
    keep = doc.add_paragraph("第二段：被保留的段落。" * 6)
    drop = doc.add_paragraph("第三段：不会被克隆的段落。" * 6)

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "7")
    start.set(qn("w:name"), "OLE_LINK7")
    keep._p.insert(0, start)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "7")
    drop._p.append(end)

    doc.save(str(path))
    return path


def test_clone_drops_bookmarks_whose_partner_was_not_cloned():
    """Word reported "unreadable content" on a bookmarkStart with no bookmarkEnd.

    Subsetting a body by paragraph index can keep one half of a paired range
    marker and drop the other. The shipped teacher edition had 8 bookmarkStart
    and 6 bookmarkEnd (OLE_LINK7/OLE_LINK8) and Word demanded a repair on open.
    """
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        src = _bookmarked(tmp / "a.docx")
        assert _count(src, "bookmarkStart") == 1 and _count(src, "bookmarkEnd") == 1

        # keep the paragraph holding bookmarkStart, drop the one holding bookmarkEnd
        out = ds.clone_subset(src, [0, 1], tmp / "c.docx")
        assert _count(out, "bookmarkStart") == 0, "the orphaned start must be removed"
        assert _count(out, "bookmarkEnd") == 0
        ds.validate(out)  # the gate that would have caught this before shipping

        # a bookmark whose two halves both survive is left alone
        both = ds.clone_subset(src, [0, 1, 2], tmp / "d.docx")
        assert _count(both, "bookmarkStart") == 1
        assert _count(both, "bookmarkEnd") == 1
        ds.validate(both)


def test_validate_rejects_an_unpaired_bookmark():
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        src = _bookmarked(tmp / "a.docx")
        out = ds.clone_subset(src, [0, 1, 2], tmp / "c.docx")

        # re-introduce the defect behind clone_subset's back
        with zipfile.ZipFile(out) as z:
            entries = [(i, z.read(i.filename)) for i in z.infolist()]
        doc = etree.fromstring(dict((i.filename, d) for i, d in entries)["word/document.xml"])
        for el in doc.iter(W + "bookmarkEnd"):
            el.getparent().remove(el)
        rewritten = etree.tostring(doc, xml_declaration=True, encoding="UTF-8", standalone=True)
        broken = tmp / "broken.docx"
        with zipfile.ZipFile(broken, "w", zipfile.ZIP_DEFLATED) as zout:
            for info, data in entries:
                zout.writestr(info, rewritten if info.filename == "word/document.xml" else data)

        try:
            ds.validate(broken)
        except RuntimeError as exc:
            assert "bookmarkStart" in str(exc)
        else:
            raise AssertionError("validate must reject an unpaired range marker")


def test_markdown_symbols_never_reach_word():
    """Word prints text verbatim, so "**原句：**" lands on the page as asterisks.

    The model is told not to emit Markdown, but the pipeline's own renderers used
    to add it — they were written for the Pandoc era and feed both sinks.
    """
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        src = _fixture(tmp / "a.docx")
        out = ds.clone_subset(src, [0], tmp / "c.docx")
        ds.decorate(out, notes=[
            (ds.NOTE_BODY, "1. **原句：** He ran.\n   - `crucial`：关键的\n# 标题"),
            (ds.NOTE_BODY, "答案填入 ____36____ 处"),
        ])

        text = db.read_docx(out).text
        assert "**" not in text
        assert "`" not in text
        assert "原句：" in text and "crucial" in text, "only the syntax goes, not the content"
        assert "标题" in text, "a heading's text stays, only the # goes"
        # underscores are NOT Markdown here: ____36____ is a fill-in blank
        assert "____36____" in text


def test_template_demo_content_does_not_ride_along():
    """The reference templates ship with a demo table as well as demo paragraphs.

    The builders only deleted `doc.paragraphs`, so a stray "Table 1 2" grid was
    printed at the top of every answer sheet that shipped.
    """
    import export_docx_splice as ex

    for name in ("student_reference.docx", "answers_reference.docx"):
        template = ex.TEMPLATE_DIR / name
        with zipfile.ZipFile(template) as z:
            source = etree.fromstring(z.read("word/document.xml"))
        assert source.findall(".//" + W + "tbl"), f"{name} is expected to carry a demo table"

        doc = ds.blank_template(template)
        body = doc.element.body
        kept = [c for c in body if isinstance(c.tag, str)]
        assert all(c.tag.endswith("}sectPr") for c in kept), "only page setup may survive"


def test_scrub_removes_lastprinted_instead_of_blanking_it():
    """An empty date is not "no date", it is a malformed date, and Word rejects it.

    scrub_metadata used to set cp:lastPrinted to "" — producing
    <cp:lastPrinted></cp:lastPrinted>, an empty string where the schema wants a
    dateTime. Word answered with 「发现无法读取的内容」. Only the student and
    teacher editions carried it (they clone the source paper's docProps); the
    answer sheet is built from a template that has none, which is exactly why it
    was the one file that never complained.
    """
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        src = _fixture(tmp / "a.docx")

        # give the fixture a printed date, like a real exam paper has
        with zipfile.ZipFile(src) as z:
            entries = [(i, z.read(i.filename)) for i in z.infolist()]
        core = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<cp:coreProperties '
            'xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
            'xmlns:dc="http://purl.org/dc/elements/1.1/" '
            'xmlns:dcterms="http://purl.org/dc/terms/" '
            'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
            "<dc:creator>某位老师</dc:creator>"
            "<cp:lastPrinted>2026-05-30T03:25:00Z</cp:lastPrinted>"
            '<dcterms:created xsi:type="dcterms:W3CDTF">2026-05-29T21:37:00Z</dcterms:created>'
            "</cp:coreProperties>"
        ).encode()
        with zipfile.ZipFile(src, "w", zipfile.ZIP_DEFLATED) as zout:
            for info, data in entries:
                zout.writestr(info, core if info.filename == "docProps/core.xml" else data)

        out = ds.clone_subset(src, [0], tmp / "c.docx")
        ds.scrub_metadata(out, "标题")

        assert ds.empty_typed_core_properties(out) == [], "an empty dateTime must never survive"
        with zipfile.ZipFile(out) as z:
            core_xml = z.read("docProps/core.xml").decode()
        assert "lastPrinted" not in core_xml, "the field is removed, not blanked"
        assert "某位老师" not in core_xml, "the original author must still be scrubbed"
        assert "2026-05-29T21:37:00Z" in core_xml, "valid dates are left alone"


def test_table_borders_land_before_tbllook():
    """CT_TblPrBase is a sequence — tblBorders after tblLook is invalid OOXML.

    python-docx's add_table() emits tblW + tblLook, so simply appending the
    borders put them last, which is out of order and makes Word repair the file.
    """
    from docx import Document
    from docx.shared import Cm

    with tempfile.TemporaryDirectory() as tmp:
        doc = Document()
        for section in doc.sections:
            section.page_width, section.page_height = Cm(21.0), Cm(29.7)
        table = doc.add_table(rows=2, cols=2)
        ds.set_table_borders(table)

        names = [c.tag.rsplit("}", 1)[-1] for c in table._tbl.tblPr]
        assert "tblBorders" in names
        ranks = [ds.TBL_PR_ORDER.index(n) for n in names if n in ds.TBL_PR_ORDER]
        assert ranks == sorted(ranks), f"tblPr children out of schema order: {names}"

        out = Path(tmp) / "t.docx"
        doc.save(str(out))
        ds.validate(out)


def test_heading_is_prepended_not_appended():
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        src = _fixture(tmp / "a.docx")
        out = ds.clone_subset(src, list(range(len(db.read_docx(src).body_children))), tmp / "c.docx")
        ds.decorate(out, heading="阅读A｜来源：某卷")
        assert db.read_docx(out).text.startswith("阅读A｜来源：某卷")


def test_answer_sheet_is_a4_even_when_the_template_is_missing():
    # The packaged app could not find its templates (PyInstaller flattens
    # scripts/, so the __file__-relative lookup pointed outside the bundle), and
    # python-docx quietly defaulted to US Letter — the export then died on its own
    # A4 check. Page size must not depend on the template being found.
    import export_docx_splice as ex

    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        original = ex.TEMPLATE_DIR
        ex.TEMPLATE_DIR = tmp / "does-not-exist"
        try:
            out = ex._answers_doc([], tmp / "answers.docx", pipeline=None)
        finally:
            ex.TEMPLATE_DIR = original

        with zipfile.ZipFile(out) as z:
            dx = etree.fromstring(z.read("word/document.xml"))
        pg = dx.find(".//" + W + "sectPr/" + W + "pgSz")
        width, height = int(pg.get(W + "w")), int(pg.get(W + "h"))
        assert 11800 <= width <= 12050 and 16700 <= height <= 17000, f"got {width}x{height}, not A4"


def test_exporter_can_find_its_templates():
    # Would have caught the packaged-app failure: the exporter resolves templates
    # through its own constant, so that is what must be checked.
    import export_docx_splice as ex
    import pdf_ingest

    assert (ex.TEMPLATE_DIR / "answers_reference.docx").exists()
    assert (ex.TEMPLATE_DIR / "student_reference.docx").exists()
    assert pdf_ingest.TEMPLATE.exists()


def test_validate_rejects_non_a4_page():
    from docx.shared import Cm

    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        src = _fixture(tmp / "a.docx")

        # python-docx defaults to US Letter; a paper for a Chinese classroom
        # must be A4, so validate() has to catch this rather than ship it.
        out = ds.clone_subset(src, [0], tmp / "letter.docx")
        try:
            ds.validate(out)
            raise AssertionError("validate() accepted a US Letter page")
        except RuntimeError as exc:
            assert "A4" in str(exc)

        d = Document(str(src))
        for section in d.sections:
            section.page_width, section.page_height = Cm(21.0), Cm(29.7)
        d.save(str(src))
        a4 = ds.clone_subset(src, [0], tmp / "a4.docx")
        ds.validate(a4)  # must not raise


def _docprops(path: Path) -> str:
    with zipfile.ZipFile(path) as z:
        return "".join(
            z.read(n).decode("utf-8", "replace") for n in z.namelist() if n.startswith("docProps")
        )


def _a4(path: Path) -> Path:
    from docx.shared import Cm

    doc = Document(str(path))
    for section in doc.sections:
        section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    doc.save(str(path))
    return path


def _sz_values(path: Path) -> set[str]:
    with zipfile.ZipFile(path) as z:
        dx = etree.fromstring(z.read("word/document.xml"))
    return {e.get(W + "val") for e in dx.findall(".//" + W + "sz")}


def test_normalize_format_pins_house_style_but_keeps_emphasis_and_media():
    """统一格式：小四 + Times New Roman/宋体 + single spacing, no snap-to-grid.

    The teacher asked for every exported file to read the same, but the questions are
    clones of a dozen source papers each in its own font and size. normalize_format
    forces the house style over the merged body while leaving the things that carry a
    question's meaning — bold, italics, the underline that draws a fill-in blank, the
    table and the image — untouched.
    """
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        src = _a4(_fixture(tmp / "a.docx"))
        out = ds.clone_subset(src, list(range(len(db.read_docx(src).body_children))), tmp / "c.docx")

        before = {tag: _count(out, tag) for tag in ("b", "u", "i", "tbl", "drawing")}
        ds.normalize_format(out)

        # emphasis, blanks, table and picture all survive
        for tag, n in before.items():
            assert _count(out, tag) == n, f"<w:{tag}> changed during normalize"
        assert ds.media_count(out) == 1
        ds.validate(out)  # still a well-formed A4 file

        # every run is 小四 (24 half-points) and nothing else
        assert _sz_values(out) == {"24"}, f"sizes are not uniformly 小四: {_sz_values(out)}"

        with zipfile.ZipFile(out) as z:
            dx = etree.fromstring(z.read("word/document.xml"))
        body = dx.find(W + "body")
        paras = body.findall(".//" + W + "p")
        # snap-to-grid is off on every paragraph, and the "same-style spacing" box
        # (contextualSpacing) is never emitted
        assert dx.findall(".//" + W + "contextualSpacing") == []
        for p in paras:
            snap = p.find(W + "pPr/" + W + "snapToGrid")
            assert snap is not None and snap.get(W + "val") == "0"
            spacing = p.find(W + "pPr/" + W + "spacing")
            assert spacing.get(W + "before") == "0" and spacing.get(W + "after") == "0"
            assert spacing.get(W + "line") == "240" and spacing.get(W + "lineRule") == "auto"
        # English in Times New Roman, 中文 in 宋体
        for fonts in dx.findall(".//" + W + "r/" + W + "rPr/" + W + "rFonts"):
            assert fonts.get(W + "ascii") == "Times New Roman"
            assert fonts.get(W + "eastAsia") == "宋体"


def test_normalize_format_size_is_configurable():
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        src = _a4(_fixture(tmp / "a.docx"))
        out = ds.clone_subset(src, list(range(len(db.read_docx(src).body_children))), tmp / "c.docx")
        ds.normalize_format(out, half_points=21)  # 五号 = 10.5pt
        assert _sz_values(out) == {"21"}


def test_scrub_metadata_removes_original_author_after_merge():
    # Must be asserted on a MERGED file, not just a clone: python-docx's
    # core_properties setter does not reach the core.xml that docxcompose
    # writes, so scrubbing a clone passed while the real export still leaked
    # the name of the teacher who wrote the source paper.
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        parts = []
        for i, name in enumerate(("a", "b")):
            src = _fixture(tmp / f"{name}.docx", marker=name.upper())
            d = Document(str(src))
            d.core_properties.author = "某位老师"
            d.core_properties.last_modified_by = "某位老师"
            d.save(str(src))
            parts.append(ds.clone_subset(src, [0, 1], tmp / f"p{i}.docx"))

        assert "某位老师" in _docprops(parts[0]), "the clone should inherit it (that's the bug)"

        merged = ds.merge(parts, tmp / "merged.docx")
        ds.scrub_metadata(merged, "学生版")

        assert "某位老师" not in _docprops(merged), "original author leaked into the export"
