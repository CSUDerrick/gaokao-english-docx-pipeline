"""The paper's own 【N题详解】 blocks, sliced per question.

The fixture reproduces the two defects the three sample papers actually have,
because both of them broke the first (character-offset) implementation:

* a paper that explains some sections and silently skips others, and
* a mistyped question number (湖北 writes 【14题详解】 where 34 belongs).

``input_docx/`` is gitignored, so this builds its own paper rather than reading
one — same reasoning as ``test_docx_splice.py``.
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import answer_explanation as ax  # noqa: E402
import docx_blocks as db  # noqa: E402
from docx import Document  # noqa: E402

# reading_a (21-23) and reading_d (32-35) are explained; reading_b/c are not —
# and 34 is mistyped as 14, exactly as 湖北 does it.
ANSWER_SECTION = [
    "参考答案",
    "21—23. CBC 24—27 DAAD 28—31 BBAC 32—35. DCDB",
    "第二部分 阅读理解",
    "A",
    "21—23. CBC",
    "【导语】本文是一篇应用文，介绍了一个度假中心的设施。",
    "【21题详解】",
    "细节理解题。根据第一段可知，该中心提供露营服务。故选C。",
    "【22题详解】",
    "推理判断题。根据第二段可知，价格随季节浮动。故选B。",
    "【23题详解】",
    "主旨大意题。全文围绕度假中心展开。故选C。",
    "D",
    "32—35. DCDB",
    "【导语】本文是一篇书评，介绍了一本关于森林的著作。",
    "【32题详解】",
    "细节理解题。根据第一段可知，作者是一名护林员。故选D。",
    "【33题详解】",
    "词义猜测题。根据上下文可知，该词意为“干预”。故选C。",
    "【14题详解】",
    "推理判断题。根据第三段可知，作者主张减少人为干预。故选D。",
    "【35题详解】",
    "观点态度题。作者对这本书持赞赏态度。故选B。",
    "第四部分 写作",
    "One possible version",
    "Dear Chris,",
    "I am writing to share with you the news that our school has launched a new plan.",
    "Yours,",
    "Li Hua",
]


def _paper(tmp: Path) -> Path:
    doc = Document()
    doc.add_paragraph("2026 届高三英语模拟卷")
    for _ in range(6):
        doc.add_paragraph("A" * 40)  # body filler, so the answer section is not at index 0
    for line in ANSWER_SECTION:
        doc.add_paragraph(line)
    path = tmp / "paper.docx"
    doc.save(str(path))
    return path


def _load(tmp: Path) -> tuple[db.DocxDoc, ax.OfficialExplanations]:
    doc = db.read_docx(_paper(tmp))
    start = doc.text.index("参考答案")
    return doc, ax.OfficialExplanations(doc, start)


def test_an_explained_section_yields_its_guide_and_every_detail_block():
    with tempfile.TemporaryDirectory() as tmp:
        doc, official = _load(Path(tmp))
        lo, hi = official.blocks_for([21, 22, 23])
        text = "\n".join(b.text for b in doc.blocks if lo <= b.body_index < hi)

        assert text.startswith("【导语】"), "the passage's 导语 introduces its questions"
        for number in (21, 22, 23):
            assert f"【{number}题详解】" in text
        assert "32题详解" not in text, "reading_d must not bleed into reading_a"
        assert "21—23. CBC" not in text, "the answer run is a group boundary, not content"


def test_a_section_the_paper_never_explains_comes_back_empty():
    with tempfile.TemporaryDirectory() as tmp:
        _, official = _load(Path(tmp))
        assert official.blocks_for([24, 25, 26, 27]) == []
        assert official.text_for([24, 25, 26, 27]) == ""


def test_a_mistyped_question_number_still_rides_along_with_its_passage():
    # 【14题详解】 is question 34. It sits between 33 and 35, so a contiguous range
    # keeps it — dropping it would lose a question's explanation entirely.
    with tempfile.TemporaryDirectory() as tmp:
        doc, official = _load(Path(tmp))
        assert official.strays == [14], "the typo is reported, not silently swallowed"

        lo, hi = official.blocks_for([32, 33, 34, 35])
        text = "\n".join(b.text for b in doc.blocks if lo <= b.body_index < hi)
        assert "【14题详解】" in text
        assert "作者主张减少人为干预" in text, "question 34's explanation survives"
        assert "【35题详解】" in text


def test_the_model_essay_is_not_mistaken_for_an_explanation():
    # The essay follows the last 详解 with no heading between them. Crediting it
    # to question 35 would print the whole 参考范文 under a reading question.
    with tempfile.TemporaryDirectory() as tmp:
        doc, official = _load(Path(tmp))
        lo, hi = official.blocks_for([32, 33, 34, 35])
        text = "\n".join(b.text for b in doc.blocks if lo <= b.body_index < hi)
        assert "Li Hua" not in text
        assert "One possible version" not in text
        assert "Dear Chris" not in text


def test_a_range_that_would_clone_another_section_is_refused():
    # Force the overlap the guard exists for: ask for 21-23 *and* 33, which spans
    # reading_a's blocks, reading_d's 导语, and 32's explanation.
    with tempfile.TemporaryDirectory() as tmp:
        _, official = _load(Path(tmp))
        assert official.blocks_for([21, 22, 23, 33]) == [], "would drag in question 32"


def test_a_paper_with_no_answer_section_explains_nothing():
    with tempfile.TemporaryDirectory() as tmp:
        doc = db.read_docx(_paper(Path(tmp)))
        official = ax.OfficialExplanations(doc, None)
        assert official.blocks_for([21, 22, 23]) == []


def test_a_writing_item_hands_the_model_the_papers_reference_essay():
    # The paper writes no 【N题详解】 for 应用文/读后续写 — its answer *is* the 参考范文,
    # which the segmenter stores as the answer key. Returning "" for those left the
    # writing prompt telling the model to beat a model essay it was never shown,
    # while the teacher edition printed that same essay directly above our answer.
    import gaokao_english_docx_pipeline as pipeline

    segment = {
        "section": "practical_writing",
        "official_explanation_blocks": [],
        "answer_key": "Dear Mr. Smith,\nI hope this email finds you well.\nYours,\nLi Hua",
    }
    assert "Li Hua" in pipeline.official_explanation_text(segment)

    # A paper that supplied no essay at all still comes back empty, so the prompt
    # correctly falls through to "原卷未提供官方解析，请你独立写出完整解析".
    segment["answer_key"] = pipeline.NO_ANSWER_MARKER
    assert pipeline.official_explanation_text(segment) == ""

    # And a reading item with no explanation blocks stays empty — its answer letters
    # already reach the model through {{ANSWER_KEY}}.
    assert pipeline.official_explanation_text(
        {"section": "reading_b", "official_explanation_blocks": [], "answer_key": [{"number": "24", "answer": "C"}]}
    ) == ""


def test_question_numbers_comes_from_the_section_not_the_answer_key():
    # A gappy answer key must not shrink the item's question set: 江苏's grammar
    # answers run together as "58. 59. to", so the key has no 59 — and asking for
    # 56-58,60-65 made 59's (correctly placed) explanation look like a foreign
    # item's, which tripped the guard and dropped the whole passage.
    gappy = [{"number": str(n), "answer": "x"} for n in (56, 57, 58, 60)]
    assert ax.question_numbers("grammar", gappy) == list(range(56, 66))
    assert ax.question_numbers("reading_b", []) == [24, 25, 26, 27]
    assert ax.question_numbers("cloze", "原卷未提供答案") == list(range(41, 56))
    assert ax.question_numbers("practical_writing", []) == [], "writing has no numbered questions"


def test_a_gappy_answer_key_still_clones_the_whole_passage():
    with tempfile.TemporaryDirectory() as tmp:
        doc, official = _load(Path(tmp))
        # The fixture's 34 is mistyped as 14, so a key built from the 详解 headings
        # would be missing 34 — the exact shape of the 江苏 bug.
        gappy = [{"number": str(n), "answer": "x"} for n in (32, 33, 35)]
        lo, hi = official.blocks_for(ax.question_numbers("reading_d", gappy))
        text = "\n".join(b.text for b in doc.blocks if lo <= b.body_index < hi)
        assert "【35题详解】" in text
