"""Regressions for the AI re-segmentation fallback.

A real run died with `RuntimeError: Segment JSON parse failed`, losing the whole
pipeline. Three separate faults were behind it, and each is pinned here:

1. A long paper pushed the reply past --segment-max-tokens, cutting the JSON
   mid-string, and the parser discarded the entire (mostly complete) response.
2. The fallback is only an attempt to *improve* a paper the local pass already
   handled, yet its failure killed the run and threw the good result away.
3. AI-segmented items carried no source_blocks, so the clone export would have
   silently dropped those questions from the Word file.
"""

from __future__ import annotations

import argparse
import importlib.util
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import docx_blocks as db  # noqa: E402

_spec = importlib.util.spec_from_file_location("pipeline", ROOT / "scripts" / "gaokao_english_docx_pipeline.py")
pipeline = importlib.util.module_from_spec(_spec)
sys.modules["pipeline"] = pipeline
_spec.loader.exec_module(pipeline)


TRUNCATED = """{
  "source_doc": "试卷.docx",
  "segments": [
    {"section": "reading_a", "display_section": "阅读A", "item_label": "21-23",
     "question_text": "A\\nUNICEF is the world leader.", "answer_key": []},
    {"section": "reading_b", "display_section": "阅读B", "item_label": "24-27",
     "question_text": "B\\nRecently I was talking.", "answer_key": []},
    {"section": "continuation_writing", "display_section": "读后续写", "item_label": "第二节",
     "question_text": "Adam came to the tryout. The players were divided into two groups and"""


def test_truncated_reply_keeps_the_segments_that_completed():
    parsed = pipeline.parse_model_json(TRUNCATED)
    assert isinstance(parsed, dict), "a cut-off reply must not be thrown away wholesale"
    assert parsed["_truncated"] is True
    labels = [s["item_label"] for s in parsed["segments"]]
    assert labels == ["21-23", "24-27"], "the two complete segments must survive"


def test_intact_json_is_untouched():
    parsed = pipeline.parse_model_json('{"source_doc": "a.docx", "segments": [{"item_label": "x"}]}')
    assert parsed["segments"] == [{"item_label": "x"}]
    assert "_truncated" not in parsed


def test_segment_max_tokens_is_large_enough_for_a_long_paper():
    # 4000 truncated a real 读后续写 unit mid-sentence.
    assert pipeline.parse_args(["input_docx"]).segment_max_tokens >= 16000


def _fixture(path: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_paragraph("A")
    doc.add_paragraph("UNICEF is the world leader in delivering vital supplies to children.")
    doc.add_paragraph("21. What does the author mean by that sentence?")
    doc.add_paragraph("B")
    doc.add_paragraph("Recently I was talking with the mother of a student I taught.")
    doc.save(str(path))
    return path


def test_ai_segment_text_is_anchored_back_to_source_blocks():
    # Without this the export cannot clone the original paragraphs, and the
    # question would vanish from the Word file.
    with tempfile.TemporaryDirectory() as tmp:
        src = _fixture(Path(tmp) / "a.docx")
        doc = db.read_docx(src)

        question = "A\nUNICEF is the world leader in delivering vital supplies to children.\n21. What does the author mean by that sentence?"
        lo, hi = db.find_block_range(doc, question)

        recovered = "\n".join(b.text for b in doc.blocks if lo <= b.body_index < hi)
        assert "UNICEF" in recovered
        assert "Recently I was talking" not in recovered, "must not swallow the next passage"


def test_unlocatable_text_reports_no_range_rather_than_guessing():
    with tempfile.TemporaryDirectory() as tmp:
        src = _fixture(Path(tmp) / "a.docx")
        doc = db.read_docx(src)
        # The model abbreviated instead of quoting; cloning a guessed range would
        # put the wrong paragraphs in the paper.
        assert db.find_block_range(doc, "略") == []
        assert db.find_block_range(doc, "") == []


def test_fallback_failure_keeps_the_local_segmentation():
    # The local pass already produced usable, cloneable segments. A failing model
    # call must not take the run down with it.
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        indir = tmp / "in"
        indir.mkdir()
        _fixture(indir / "paper.docx")

        args = pipeline.parse_args([str(indir), "--out", str(tmp / "out"), "--segment-input", "local"])
        args.api_key = "test-key"

        original = pipeline.segment_docx_file
        calls = {"n": 0}

        def flaky(docx, a, out_dir):
            # rough mode blows up the way the real run did; local mode still works
            if getattr(a, "segment_input", "local") == "rough":
                calls["n"] += 1
                raise RuntimeError("Segment JSON parse failed for paper.docx / 读后续写")
            return original(docx, a, out_dir)

        forced = pipeline.evaluate_document

        def always_warn(doc, rows):
            result = forced(doc, rows)
            result["needs_model_fallback"] = True
            result["grade"] = "WARN"
            result["structural_issues"] = ["forced"]
            return result

        pipeline.segment_docx_file = flaky
        pipeline.evaluate_document = always_warn
        try:
            rows = pipeline.run_segment(args)  # must NOT raise
        finally:
            pipeline.segment_docx_file = original
            pipeline.evaluate_document = forced

        assert calls["n"] >= 1, "the fallback should have been attempted"
        assert rows, "the local segmentation must survive the failed fallback"
        assert all(r.get("segment_path") for r in rows)
